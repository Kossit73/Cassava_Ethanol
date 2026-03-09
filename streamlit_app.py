"""Streamlit dashboard for the Cassava bioethanol financial model."""

from __future__ import annotations

import copy
import io
import json
import re
import tempfile
from datetime import datetime
import textwrap
from collections import OrderedDict
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import numpy as np
import pandas as pd
import plotly.express as px
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
import streamlit as st

from bioethanol_model import CassavaBioethanolModel
from bioethanol_model.exporter import export_to_excel
from bioethanol_model.increment import apply_yearly_increment
from bioethanol_model.inputs import (
    EditableTable,
    InputLandingPage,
    ProjectionHorizon,
    default_input_page,
)
from bioethanol_model.scenario import (
    ScenarioConfig,
    credit_committee_scenario_configs,
    goal_seek_to_target,
    reverse_stress_test,
    scenario_comparison,
    scenario_parameter_catalog,
)
from bioethanol_model.schedules import (
    ANIMAL_FEED_TON_PER_TON,
    ETHANOL_LITRES_PER_TON,
    ExpenseSummary,
    WorkingCapitalOutput,
    compute_production_tables,
    compute_staff_schedule,
)
from bioethanol_model.sensitivity import (
    DEFAULT_MONTE_CARLO_ITERATIONS,
    DEFAULT_MONTE_CARLO_SEED,
    MONTE_CARLO_DISTRIBUTIONS,
    MONTE_CARLO_PARAMETER_ADAPTERS,
    MONTE_CARLO_PARAMETER_COLUMNS,
    MONTE_CARLO_TEXT_COLUMNS,
    SCENARIO_PARAMETER_NAMES,
    SensitivityScenario,
    available_monte_carlo_distributions,
    default_monte_carlo_parameters,
    monte_carlo_simulation,
    run_sensitivity,
    tornado_chart_inputs,
)


st.set_page_config(page_title="Cassava Bioethanol Model", layout="wide")

MODEL_VERSION_KEY = "model_version"
MC_CACHE_KEY = "mc_cache_store"
SENSITIVITY_CACHE_KEY = "sensitivity_cache"
SCENARIO_CACHE_KEY = "scenario_cache"

# Columns that are derived from other inputs and should not be editable via the
# "Modify Default Inputs & Figures" pane or the general data editor. The map is
# keyed by landing-page table name to keep the behaviour scoped and explicit.
DERIVED_COLUMN_MAP = {
    "Production Monthly": {"Ethanol litres", "Animal Feed ton"},
    "Production Annual": {"Ethanol litres", "Animal Feed ton"},
}


# Predefined category options surfaced in the "Modify Default Inputs & Figures"
# editor. Users can still supply custom values by selecting the explicit custom
# option exposed by the editor for each table.
DIRECT_COST_CATEGORY_OPTIONS = [
    "Cassava Feedstock",
    "Enzymes & Chemicals",
    "Energy Cost",
]

CATEGORY_SELECT_OPTIONS = {
    ("Direct Costs Monthly", "Cost Category"): {
        "options": DIRECT_COST_CATEGORY_OPTIONS,
        "allow_custom": False,
    },
    (
        "Other Opex Monthly",
        "Category",
    ): {
        "options": [
            "Service Contracts",
            "General Administration",
            "Research & Development",
            "Energy Cost",
            "Sales & Marketing",
        ],
        "allow_custom": True,
    },
    (
        "Accounts Receivable & Other Assets",
        "Metric",
    ): {
        "options": [
            "Receivables days",
            "Inventory days",
            "Prepaid expense days",
            "Other assets percent of revenue",
        ],
        "allow_custom": True,
    },
    (
        "Accounts Payable",
        "Metric",
    ): {
        "options": [
            "Payables days",
            "Other payable days",
        ],
        "allow_custom": True,
    },
}

YEARLY_INCREMENT_CONFIG = {
    "Production Monthly": {
        "date_column": "Start Month",
        "frequency": "M",
        "value_columns": ["Cassava ton"],
        "match_columns": [],
        "description": "Apply an annual growth/decline to cassava tonnage.",
    },
    "Production Annual": {
        "date_column": "Year",
        "frequency": "Y",
        "value_columns": ["Cassava ton"],
        "match_columns": [],
        "description": "Scale annual cassava tonnage by a yearly percentage.",
    },
    "Direct Costs Monthly": {
        "date_column": "Month",
        "frequency": "M",
        "value_columns": ["Amount"],
        "match_columns": ["Cost Category"],
        "description": "Increase or decrease the selected cost category each year.",
    },
    "Staff Costs Monthly": {
        "date_column": "Month",
        "frequency": "M",
        "value_columns": ["Headcount", "Cost"],
        "match_columns": ["Department"],
        "description": "Apply annual changes to staffing levels or department spend.",
    },
    "Other Opex Monthly": {
        "date_column": "Month",
        "frequency": "M",
        "value_columns": ["Amount"],
        "match_columns": ["Category"],
        "description": "Cascade the yearly percentage to this operating expense.",
    },
    "Accounts Receivable & Other Assets": {
        "date_column": "Effective Month",
        "frequency": "M",
        "value_columns": ["Value"],
        "match_columns": ["Metric"],
        "description": "Roll forward the working-capital metric with an annual rate.",
    },
    "Accounts Payable": {
        "date_column": "Effective Month",
        "frequency": "M",
        "value_columns": ["Value"],
        "match_columns": ["Metric"],
        "description": "Apply the yearly percentage to the payable policy value.",
    },
    "Inflation Schedule": {
        "date_column": "Year",
        "frequency": "Y",
        "value_columns": ["CPI", "FX Index", "Tariff Escalation"],
        "match_columns": [],
        "description": "Project inflation indices with an annual growth rate.",
    },
    "Loan Schedule": {
        "date_column": "Start Month",
        "frequency": "M",
        "value_columns": ["Loan Amount"],
        "match_columns": ["Loan"],
        "description": "Escalate the selected facility amount year over year.",
    },
}


def _increment_horizon_value(config: Dict[str, object], projection: ProjectionHorizon) -> object | None:
    """Return the terminal period for yearly increments based on *projection*."""

    frequency = str(config.get("frequency", "")).upper()
    if frequency.startswith("M"):
        return f"{projection.end_year:04d}-12"
    if frequency.startswith("Y"):
        return projection.end_year
    return None


CHANGE_BUTTON_CONFIG = {
    "Production Monthly": {
        "type": "month",
        "start_year_offset": 1,
        "help": (
            "Select the month where the new production plan should begin, click "
            "**Add change**, and then edit the newly inserted row. The model will "
            "cascade the cassava tonnage (and the derived ethanol and animal feed "
            "outputs) to all later months automatically."
        ),
    },
    "Direct Costs Monthly": {
        "type": "month",
        "start_year_offset": 1,
        "help": (
            "Pick the month the revised cost takes effect and use **Add change** to "
            "insert a row you can edit. The schedule keeps the change for that month "
            "and beyond."
        ),
    },
    "Staff Costs Monthly": {
        "type": "month",
        "start_year_offset": 1,
        "help": (
            "Choose the effective month for the staffing update, press **Add change**, "
            "and adjust the new row. Future months inherit the change unless another "
            "override is added."
        ),
    },
    "Other Opex Monthly": {
        "type": "month",
        "start_year_offset": 1,
        "help": (
            "Select the start month for the new opex figure, click **Add change**, and "
            "enter the updated values on the inserted row to carry them forward."
        ),
    },
    "Accounts Receivable & Other Assets": {
        "type": "month",
        "start_year_offset": 1,
        "help": (
            "Use **Add change** after picking the effective month to insert a new "
            "policy row. The working-capital calculations will apply it from that "
            "point onward."
        ),
    },
    "Accounts Payable": {
        "type": "month",
        "start_year_offset": 1,
        "help": (
            "Insert a new policy row by selecting the effective month and clicking "
            "**Add change**; the updated payable settings will roll forward "
            "automatically."
        ),
    },
    "Loan Schedule": {
        "type": "month",
        "start_year_offset": 0,
        "help": (
            "Choose the draw month for a facility and press **Add change** to add a "
            "new loan row tied to that date."
        ),
    },
    "Inflation Schedule": {
        "type": "year",
        "start_year_offset": 0,
        "help": (
            "Select the year that a new index value applies to, click **Add change**, "
            "and edit the inserted row."
        ),
    },
}

DEFAULT_SENSITIVITY_SCENARIOS: List[SensitivityScenario] = [
    SensitivityScenario("Corporate tax +1pp", "Corporate tax rate", 0.01),
    SensitivityScenario("Corporate tax -1pp", "Corporate tax rate", -0.01),
    SensitivityScenario("Discount rate +1pp", "Discount rate", 0.01),
    SensitivityScenario("Discount rate -1pp", "Discount rate", -0.01),
]

DEFAULT_SCENARIO_CONFIGS: List[ScenarioConfig] = [
    ScenarioConfig("Higher investor share", {"Investor share capital": 0.55}),
    ScenarioConfig("Lower investor share", {"Investor share capital": 0.35}),
    ScenarioConfig("Lower tax", {"Corporate tax rate": 0.24}),
]

TORNADO_DRIVERS: List[Tuple[str, float]] = [
    ("Corporate tax rate", 1.0),
    ("Investor share capital", 1.0),
    ("Owner share capital", 1.0),
    ("Discount rate", 1.0),
]

MC_PARAMETER_STATE_KEY = "mc_parameter_config"
MC_SELECTED_PARAMETER_STATE_KEY = "mc_selected_parameters"
MC_ITERATION_STATE_KEY = "mc_iteration_setting"
MC_SEED_STATE_KEY = "mc_random_seed"

SCENARIO_DEFINITIONS_KEY = "scenario_definitions"
SCENARIO_SELECTION_STATE_KEY = "scenario_builder_selection"
SCENARIO_VALUE_STATE_KEY = "scenario_builder_values"
SCENARIO_NAME_STATE_KEY = "scenario_builder_name"
SCENARIO_CLEAR_NAME_FLAG = "scenario_builder_clear_flag"

PRODUCTION_EDIT_FLAG = "production_user_edit_flag"


def _trigger_rerun() -> None:
    """Request Streamlit to rerun the script if the API is available."""

    rerun = getattr(st, "rerun", None)
    if rerun is None:
        rerun = getattr(st, "experimental_rerun", None)
    if rerun is not None:
        rerun()


def _normalise_json_value(value: object) -> object:
    if isinstance(value, float):
        if not np.isfinite(value):
            return None
        return float(value)
    return value


def _scenario_definition_signature(definitions: Iterable[Dict[str, object]]) -> str:
    normalised: List[Dict[str, object]] = []
    for entry in definitions:
        normalised.append(
            {
                "name": entry.get("name"),
                "overrides": {
                    key: _normalise_json_value(val) for key, val in entry.get("overrides", {}).items()
                },
                "deltas": {
                    key: _normalise_json_value(val) for key, val in entry.get("deltas", {}).items()
                },
            }
        )
    return json.dumps(normalised, sort_keys=True)


def _load_session_inputs() -> InputLandingPage:
    """Return the mutable input landing page stored in session state."""
    if "input_page" not in st.session_state:
        st.session_state.input_page = default_input_page()
    return st.session_state.input_page


def _mark_inputs_dirty() -> None:
    """Flag the session so financial outputs are refreshed on the next run."""

    st.session_state.inputs_dirty = True


def _table_widget_key(table: EditableTable) -> str:
    """Return the Streamlit widget key for the data editor bound to *table*."""

    return f"table_editor_{table.name.replace(' ', '_').lower()}"


def _table_editor_state_key(table: EditableTable) -> str:
    """Return the session-state cache key used to mirror the table data."""

    return f"table_cache_{table.name.replace(' ', '_').lower()}"


def _sync_table_editors(page: InputLandingPage) -> None:
    """Keep Streamlit's editor cache aligned with the latest table data."""

    for table in page.tables().values():
        key = _table_editor_state_key(table)
        table_copy = table.data.copy()
        if key not in st.session_state:
            st.session_state[key] = table_copy
        else:
            editor_value = st.session_state[key]
            if not isinstance(editor_value, pd.DataFrame) or not editor_value.equals(table_copy):
                st.session_state[key] = table_copy


def _update_table_editor_state(table: EditableTable) -> None:
    """Force cached editor data to reflect *table*'s current values."""

    df_copy = table.data.copy()
    st.session_state[_table_editor_state_key(table)] = df_copy


def _reset_table_widget(table: EditableTable) -> None:
    """Remove any bound widget state so new data is rendered on the next run."""

    widget_key = _table_widget_key(table)
    if widget_key in st.session_state:
        del st.session_state[widget_key]


def _ensure_monte_carlo_state() -> None:
    if MC_PARAMETER_STATE_KEY not in st.session_state or not isinstance(
        st.session_state[MC_PARAMETER_STATE_KEY], pd.DataFrame
    ):
        st.session_state[MC_PARAMETER_STATE_KEY] = default_monte_carlo_parameters()

    df = st.session_state[MC_PARAMETER_STATE_KEY]
    missing = [col for col in MONTE_CARLO_PARAMETER_COLUMNS if col not in df.columns]
    if missing:
        for column in missing:
            if column in MONTE_CARLO_TEXT_COLUMNS:
                df[column] = ""
            else:
                df[column] = np.nan

    df = df[list(MONTE_CARLO_PARAMETER_COLUMNS)]

    for column in MONTE_CARLO_TEXT_COLUMNS:
        df[column] = df[column].astype("string").fillna("").astype(object)

    st.session_state[MC_PARAMETER_STATE_KEY] = df

    if MC_ITERATION_STATE_KEY not in st.session_state:
        st.session_state[MC_ITERATION_STATE_KEY] = DEFAULT_MONTE_CARLO_ITERATIONS
    if MC_SEED_STATE_KEY not in st.session_state:
        st.session_state[MC_SEED_STATE_KEY] = DEFAULT_MONTE_CARLO_SEED


def _monte_carlo_parameters() -> pd.DataFrame:
    _ensure_monte_carlo_state()
    return st.session_state[MC_PARAMETER_STATE_KEY].copy()


def _monte_carlo_signature(config: pd.DataFrame, iterations: int, seed: int) -> str:
    filtered = (
        config.replace("", np.nan)
        .dropna(subset=["Parameter", "Distribution"], how="any")
        .fillna("<nan>")
    )
    ordered = filtered.sort_values(["Parameter", "Distribution"]).reset_index(drop=True)
    payload = {
        "iterations": int(iterations),
        "seed": int(seed),
        "config": ordered.to_dict(orient="records"),
    }
    return json.dumps(payload, sort_keys=True)


def _monte_carlo_parameter_library(page: InputLandingPage | None) -> pd.DataFrame:
    """Return the Monte Carlo parameter catalog derived from landing-page inputs."""

    rows: List[Dict[str, object]] = []
    for name, adapter in MONTE_CARLO_PARAMETER_ADAPTERS.items():
        base_value = np.nan
        units = adapter.units
        if isinstance(page, InputLandingPage):
            try:
                state = adapter.capture(page)
            except AttributeError:
                state = None
            if state is not None:
                base_value = state.base_value
        rows.append({"Parameter": name, "Base Value": base_value, "Units": units})

    return pd.DataFrame(rows, columns=["Parameter", "Base Value", "Units"])


def _monte_carlo_parameter_options(page: InputLandingPage | None) -> List[str]:
    """Return the ordered list of Monte Carlo parameters sourced from inputs."""

    library = _monte_carlo_parameter_library(page)
    if library.empty:
        return []
    return library["Parameter"].tolist()


def _monte_carlo_distribution_table() -> pd.DataFrame:
    """Return a help table describing supported distributions and parameters."""

    rows: List[Dict[str, str]] = []
    for name, spec in MONTE_CARLO_DISTRIBUTIONS.items():
        parts: List[str] = []
        if spec.shape_params:
            parts.append(", ".join(spec.shape_params))
        if spec.keyword_params:
            parts.append(", ".join(spec.keyword_params))
        parameter_list = ", ".join(parts) if parts else "None"
        rows.append({"Distribution": name, "Required Parameters": parameter_list})
    return pd.DataFrame(rows)


def _projection_month_options(projection: ProjectionHorizon, start_year_offset: int = 0) -> List[str]:
    start_year = projection.start_year + start_year_offset
    if start_year > projection.end_year:
        return []
    start_period = pd.Period(f"{start_year:04d}-01", freq="M")
    end_period = pd.Period(f"{projection.end_year:04d}-12", freq="M")
    if start_period > end_period:
        return []
    periods = pd.period_range(start_period, end_period, freq="M")
    return [p.strftime("%Y-%m") for p in periods]


def _projection_year_options(projection: ProjectionHorizon, start_year_offset: int = 0) -> List[int]:
    start_year = projection.start_year + start_year_offset
    if start_year > projection.end_year:
        return []
    return list(range(start_year, projection.end_year + 1))


def _apply_change_row(table: EditableTable, column: str, value: object) -> None:
    df = table.data.copy()
    if not df.empty:
        base_row = df.iloc[-1].to_dict()
    else:
        base_row = {col: None for col in table.columns}

    if column == "Year":
        base_row[column] = int(value)
    else:
        base_row[column] = str(value)

    if column in df.columns and not df.empty:
        mask = df[column].astype(str) == str(base_row[column])
    else:
        mask = pd.Series(False, index=df.index)

    if mask.any():
        idx = mask[mask].index[-1]
        for col in table.columns:
            df.at[idx, col] = base_row.get(col)
    else:
        df = pd.concat([df, pd.DataFrame([base_row])], ignore_index=True)

    if column == "Year" and "Year" in df.columns:
        df["Year"] = pd.to_numeric(df["Year"], errors="coerce")
        df = df.sort_values("Year", na_position="last").reset_index(drop=True)
    elif column in df.columns:
        try:
            order = pd.to_datetime(df[column].astype(str), errors="coerce").dt.to_period("M").dt.to_timestamp()
            df = df.assign(_order=order).sort_values("_order").drop(columns="_order").reset_index(drop=True)
        except Exception:  # pragma: no cover - defensive sorting guard
            df = df.reset_index(drop=True)

    table.set_data(df[table.columns], mark_user_input=True)
    _update_table_editor_state(table)
    _mark_inputs_dirty()


def _render_change_controls(page: InputLandingPage, table: EditableTable) -> bool:
    config = CHANGE_BUTTON_CONFIG.get(table.name)
    if not config:
        return False

    change_type = config["type"]
    offset = config.get("start_year_offset", 0)
    safe_key = table.name.replace(" ", "_").lower()

    if change_type == "month":
        month_col = _get_month_column(table.data)
        if not month_col:
            return False
        month_options = _projection_month_options(page.projection, offset)
        if not month_options:
            return False
        default_month = month_options[0]
        planning_start = page.projection.planning_start
        if planning_start and planning_start in month_options:
            default_month = planning_start
        selected_month = st.selectbox(
            "Change effective month",
            month_options,
            index=month_options.index(default_month) if default_month in month_options else 0,
            key=f"change_month_select_{safe_key}",
        )
        help_text = config.get("help")
        if help_text:
            st.caption(help_text)
        if st.button("Add change", key=f"change_month_btn_{safe_key}"):
            _apply_change_row(table, month_col, selected_month)
            return True
    elif change_type == "year":
        if "Year" not in table.columns:
            return False
        year_options = _projection_year_options(page.projection, offset)
        if not year_options:
            return False
        existing_years = pd.to_numeric(table.data.get("Year"), errors="coerce").dropna().astype(int)
        default_year = year_options[0]
        if not existing_years.empty:
            candidate = existing_years.max() + 1
            if candidate in year_options:
                default_year = candidate
        selected_year = st.selectbox(
            "Change effective year",
            year_options,
            index=year_options.index(int(default_year)) if default_year in year_options else 0,
            key=f"change_year_select_{safe_key}",
        )
        help_text = config.get("help")
        if help_text:
            st.caption(help_text)
        if st.button("Add change", key=f"change_year_btn_{safe_key}"):
            _apply_change_row(table, "Year", int(selected_year))
            return True

    return False


def _current_model_version() -> int:
    return int(st.session_state.get(MODEL_VERSION_KEY, 0))


def _bump_model_version() -> None:
    st.session_state[MODEL_VERSION_KEY] = _current_model_version() + 1


def _get_month_column(df: pd.DataFrame) -> str | None:
    """Return the preferred month column present in *df* (if any)."""

    for column in ("Effective Month", "Start Month", "Month"):
        if column in df.columns:
            return column
    return None


def _build_model_snapshot(page: InputLandingPage) -> tuple[CassavaBioethanolModel, Dict[str, object]]:
    """Create a model/result pair from a deep copy of the landing-page inputs."""

    snapshot = copy.deepcopy(page)
    model = CassavaBioethanolModel(snapshot)
    return model, model.build()


def _generate_excel_bytes(
    model: CassavaBioethanolModel, results: Dict[str, object], scenario: str
) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_path = Path(tmpdir) / "Cassava_Bioethanol_Financial_Model.xlsx"
        export_to_excel(model, temp_path, results=results, scenario=scenario)
        return temp_path.read_bytes()


def _ensure_scenario_payload(
    scenario: str, snapshot: InputLandingPage
) -> Tuple[CassavaBioethanolModel, Dict[str, object]]:
    """Return (model, results) for a scenario, computing it lazily if needed."""

    payloads: Dict[str, Tuple[CassavaBioethanolModel, Dict[str, object]]] = (
        st.session_state.setdefault("scenario_payloads", {})
    )
    if scenario not in payloads:
        with st.spinner(f"Running {scenario.replace('_', ' ').title()} scenario..."):
            model = CassavaBioethanolModel(copy.deepcopy(snapshot))
            results = model.build(scenario)
        payloads[scenario] = (model, results)
        st.session_state.scenario_payloads = payloads
    return payloads[scenario]


def _update_projection(page: InputLandingPage) -> None:
    """Render projection horizon controls within the main layout."""

    st.subheader("Projection Horizon")
    start_col, end_col = st.columns(2)
    start = start_col.number_input(
        "Start Year",
        min_value=2000,
        max_value=2100,
        value=int(page.projection.start_year),
        step=1,
        key="projection_start_year",
    )
    end = end_col.number_input(
        "End Year",
        min_value=start,
        max_value=2125,
        value=int(page.projection.end_year),
        step=1,
        key="projection_end_year",
    )
    page.projection.start_year = int(start)
    page.projection.end_year = int(end)

    month_options = pd.period_range(f"{int(start):04d}-01", f"{int(end):04d}-12", freq="M")
    month_labels = [p.strftime("%Y-%m") for p in month_options]
    default_plan = page.projection.planning_start or month_labels[0]
    if default_plan not in month_labels:
        default_plan = month_labels[0]
    planning_selection = st.selectbox(
        "Planning Start Month",
        options=month_labels,
        index=month_labels.index(default_plan) if default_plan in month_labels else 0,
        key="projection_planning_start",
    )
    page.projection.planning_start = planning_selection
    page.projection.clamp_planning_start()


def _shift_month_column(table: EditableTable, year_delta: int) -> bool:
    """Shift a table's ``Month`` column by *year_delta* years.

    When the projection horizon changes we want the editable landing-page
    tables (production, staff costs, other opex) to reflect the new start
    year.  Shifting by the delta preserves the relative spacing and keeps any
    duplicate months (e.g. multiple departments in the same month) intact.
    Returns ``True`` when an update was applied.
    """

    month_col = _get_month_column(table.data)
    if year_delta == 0 or table.data.empty or month_col is None:
        return False

    try:
        periods = pd.PeriodIndex(table.data[month_col].astype(str), freq="M")
    except Exception:  # pragma: no cover - defensive parsing guard
        return False

    shifted = periods + year_delta * 12
    new_values = shifted.astype(str)
    current_values = table.data[month_col].astype(str).reset_index(drop=True)

    if current_values.equals(pd.Series(new_values)):
        return False

    table.data.loc[:, month_col] = new_values
    return True


def _sync_projection_from_session(page: InputLandingPage) -> None:
    """Ensure the landing-page projection matches the latest widget state."""

    start_key = "projection_start_year"
    end_key = "projection_end_year"
    planning_key = "projection_planning_start"

    previous_start = int(page.projection.start_year)
    previous_end = int(page.projection.end_year)
    previous_plan = str(page.projection.planning_start)

    start = int(st.session_state.get(start_key, previous_start))
    end = int(st.session_state.get(end_key, previous_end))
    if end < start:
        end = start

    page.projection.start_year = start
    page.projection.end_year = end
    st.session_state[start_key] = start
    st.session_state[end_key] = end

    month_options = pd.period_range(f"{start:04d}-01", f"{end:04d}-12", freq="M").strftime("%Y-%m").tolist()
    plan_value = str(st.session_state.get(planning_key, previous_plan))
    if plan_value not in month_options and month_options:
        plan_value = month_options[0]
    page.projection.planning_start = plan_value
    page.projection.clamp_planning_start()
    st.session_state[planning_key] = page.projection.planning_start

    year_delta = start - previous_start
    tables_to_shift = [
        page.production_monthly,
        page.direct_costs_monthly,
        page.staff_costs_monthly,
        page.other_opex_monthly,
        page.accounts_receivable,
        page.inventory_payable,
    ]

    any_shifted = False
    for tbl in tables_to_shift:
        if _shift_month_column(tbl, year_delta):
            any_shifted = True

    if (
        any_shifted
        or start != previous_start
        or end != previous_end
        or page.projection.planning_start != previous_plan
    ):
        _mark_inputs_dirty()


def _auto_compound_production(page: InputLandingPage) -> None:
    """Cascade monthly production volumes using the growth assumptions."""

    df = page.production_monthly.data.copy()
    month_col = _get_month_column(df)
    if df.empty or month_col is None:
        return

    month_series = pd.to_datetime(df[month_col].astype(str), errors="coerce")
    valid_mask = month_series.notna()
    if not valid_mask.any():
        return

    df = df.loc[valid_mask].reset_index(drop=True)
    month_periods = month_series[valid_mask].dt.to_period("M")
    df.loc[:, month_col] = month_periods.astype(str)
    month_index = month_periods

    sort_order = np.argsort(month_index.astype(str))
    if len(sort_order) and not np.all(sort_order == np.arange(len(sort_order))):
        month_index = month_index[sort_order]
        df = df.iloc[sort_order].reset_index(drop=True)

    if getattr(month_index, "has_duplicates", False) and month_index.has_duplicates:
        keep_mask = ~month_index.duplicated(keep="last")
        if keep_mask.ndim:
            keep_indices = np.flatnonzero(keep_mask)
            month_index = month_index[keep_indices]
            df = df.iloc[keep_indices].reset_index(drop=True)
        else:  # pragma: no cover - defensive fallback for scalar mask
            month_index = month_index.unique()
            df = df.iloc[: len(month_index)].reset_index(drop=True)
        df.loc[:, month_col] = month_index.astype(str)

    month_index_set = set(month_index)

    previous_cache = st.session_state.get("production_compound_cache")
    previous_series = None
    if isinstance(previous_cache, pd.DataFrame) and not previous_cache.empty:
        prev_month_col = _get_month_column(previous_cache)
        if prev_month_col and "Cassava ton" in previous_cache.columns:
            try:
                prev_index = pd.PeriodIndex(previous_cache[prev_month_col].astype(str), freq="M")
                previous_series = pd.Series(
                    pd.to_numeric(previous_cache["Cassava ton"], errors="coerce"),
                    index=prev_index,
                )
            except Exception:  # pragma: no cover - defensive parsing guard
                previous_series = None

    numeric_columns = [
        col for col in ("Cassava ton", "Ethanol litres", "Animal Feed ton") if col in df.columns
    ]
    manual_columns = [col for col in ("Cassava ton",) if col in df.columns]
    if not numeric_columns:
        return
    if not manual_columns:
        manual_columns = numeric_columns[:1]

    growth_col = next((c for c in df.columns if "growth" in c.lower()), None)
    growth_values = pd.Series(dtype=float)
    if growth_col:
        growth_values = pd.to_numeric(df[growth_col], errors="coerce")
        growth_values.index = month_index
        if not growth_values.dropna().empty and growth_values.nunique(dropna=False) <= 1:
            base_growth = float(growth_values.dropna().iloc[0]) if not growth_values.dropna().empty else 0.0
            growth_values = pd.Series(
                [base_growth] + [np.nan] * (len(growth_values) - 1), index=month_index
            )
    else:
        growth_values = pd.Series(index=month_index, dtype=float)

    manual_periods: set[pd.Period] = set()
    planning_period: pd.Period | None = None
    first_period: pd.Period | None = None
    manual_state_key = "production_manual_periods"
    session_manual = st.session_state.get(manual_state_key, set())
    if not isinstance(session_manual, set):
        session_manual = set(session_manual)
    manual_session_periods: set[pd.Period] = set()
    for value in session_manual:
        try:
            manual_session_periods.add(pd.Period(str(value), freq="M"))
        except Exception:  # pragma: no cover - defensive parsing guard
            continue
    if getattr(page.projection, "planning_start", None):
        try:
            planning_period = pd.Period(page.projection.planning_start, freq="M")
        except Exception:  # pragma: no cover - defensive parsing guard
            planning_period = None

    if len(month_index) > 0:
        for candidate in month_index:
            if planning_period is not None and candidate < planning_period:
                continue
            first_period = candidate
            break
        if first_period is None:
            first_period = month_index[0]

    if not growth_values.empty:
        # Treat the first row as the anchor growth rate. Any subsequent entries
        # that simply repeat this base value are considered placeholders and can
        # be overridden by the compounded series. If a user enters a different
        # growth figure for a later month we keep it so that a new cascade can
        # start from that point.
        base_period = growth_values.index[0]
        base_growth = float(growth_values.iloc[0]) if pd.notna(growth_values.iloc[0]) else 0.0
        growth_values.iloc[0] = base_growth
        tolerance = 1e-9
        for idx in growth_values.index[1:]:
            raw_val = growth_values.at[idx]
            if isinstance(raw_val, (pd.Series, np.ndarray, list, tuple)):
                if len(raw_val) == 0:
                    val = np.nan
                else:
                    val = raw_val[0]
            else:
                val = raw_val

            try:
                is_missing = pd.isna(val)
            except ValueError:
                # pandas can raise when the scalar resolution is ambiguous;
                # treat it as missing so the cascade can overwrite it.
                is_missing = True

            if is_missing:
                continue

            try:
                numeric_val = float(val)
            except (TypeError, ValueError):
                continue

            if abs(numeric_val - base_growth) < tolerance:
                growth_values.at[idx] = np.nan
        manual_periods.add(base_period)
        if planning_period is not None and base_period < planning_period:
            manual_periods.discard(base_period)

    growth_periods: set[pd.Period] = set()
    for period, val in growth_values.dropna().items():
        if val is not None:
            growth_periods.add(period)

    cassava_current = pd.Series(dtype=float)
    if "Cassava ton" in df.columns:
        cassava_current = pd.to_numeric(df["Cassava ton"], errors="coerce")
        cassava_current.index = month_index

    changed_periods: set[pd.Period] = set()
    if previous_series is None:
        if first_period is not None:
            changed_periods.add(first_period)
    else:
        aligned_prev = previous_series.reindex(month_index)
        for period, value in cassava_current.items():
            prev_val = aligned_prev.get(period)

            if isinstance(prev_val, (pd.Series, np.ndarray, list, tuple)):
                prev_series = pd.Series(prev_val).dropna()
                prev_val = prev_series.iloc[0] if not prev_series.empty else np.nan

            try:
                current_missing = pd.isna(value)
            except ValueError:
                current_missing = True

            try:
                prev_missing = pd.isna(prev_val)
            except ValueError:
                prev_missing = True

            if current_missing:
                if not prev_missing:
                    changed_periods.add(period)
                continue

            try:
                numeric_val = float(value)
            except (TypeError, ValueError):
                continue

            if prev_missing or not np.isfinite(prev_val) or not np.isclose(float(prev_val), numeric_val, atol=1e-9):
                changed_periods.add(period)

    baseline_changed = False
    if first_period is not None:
        if previous_series is None:
            baseline_changed = True
        elif first_period in changed_periods:
            baseline_changed = True

    if baseline_changed:
        preserve = {p for p in changed_periods if p != first_period}
        manual_periods = set(preserve)
    else:
        manual_periods = set(manual_session_periods)
        manual_periods.update(changed_periods)

    manual_periods.update(growth_periods)

    if planning_period is not None:
        manual_periods = {p for p in manual_periods if p >= planning_period}

    if first_period is not None:
        manual_periods.add(first_period)

    manual_periods &= month_index_set

    st.session_state[manual_state_key] = {
        period.strftime("%Y-%m") for period in manual_periods
    }

    seed_df = df.copy()
    manual_mask = month_index.isin(list(manual_periods))
    for col in manual_columns:
        seed_df.loc[~manual_mask, col] = np.nan
    if growth_col:
        seed_df[growth_col] = growth_values.values

    production = compute_production_tables(
        seed_df,
        page.projection.start_year,
        page.projection.end_year,
        planning_start=page.projection.planning_start_timestamp,
    )

    monthly = production.monthly.copy()
    if monthly.empty:
        return

    monthly_reset = monthly.reset_index()
    monthly_reset["Month"] = monthly_reset["Month"].dt.to_period("M").astype(str)
    if month_col != "Month":
        monthly_reset[month_col] = monthly_reset["Month"]
        monthly_reset = monthly_reset.drop(columns=["Month"])
    if growth_col:
        display_growth = growth_values.copy()
        if isinstance(display_growth.index, pd.PeriodIndex):
            display_growth.index = display_growth.index.to_timestamp()
        if not display_growth.index.is_unique:
            display_growth = display_growth[~display_growth.index.duplicated(keep="last")]
        display_growth = display_growth.reindex(monthly.index).ffill()
        monthly_reset[growth_col] = display_growth.values

    monthly_order = [col for col in df.columns if col in monthly_reset.columns]
    new_monthly = monthly_reset[monthly_order].copy()
    for col in numeric_columns:
        if col in new_monthly.columns:
            new_monthly[col] = pd.to_numeric(new_monthly[col], errors="coerce").round(6)
    if growth_col and growth_col in new_monthly.columns:
        new_monthly[growth_col] = pd.to_numeric(new_monthly[growth_col], errors="coerce")

    current_monthly = df[monthly_order].copy()
    for col in numeric_columns:
        if col in current_monthly.columns:
            current_monthly[col] = pd.to_numeric(current_monthly[col], errors="coerce").round(6)
    if growth_col and growth_col in current_monthly.columns:
        current_monthly[growth_col] = pd.to_numeric(current_monthly[growth_col], errors="coerce")

    updated = False
    user_driven_change = previous_series is not None and bool(changed_periods)
    user_edit_flag = bool(st.session_state.get(PRODUCTION_EDIT_FLAG, False))

    if not new_monthly.equals(current_monthly):
        mark_user = user_edit_flag or user_driven_change or not page.production_monthly.placeholder
        page.production_monthly.set_data(new_monthly, mark_user_input=mark_user)
        _update_table_editor_state(page.production_monthly)
        _reset_table_widget(page.production_monthly)
        updated = True

    st.session_state["production_compound_cache"] = new_monthly.copy()

    annual = production.annual.copy()
    annual.index.name = "Year"
    annual_reset = annual.reset_index()
    annual_order = [col for col in page.production_annual.data.columns if col in annual_reset.columns]
    new_annual = annual_reset[annual_order].copy()
    for col in annual_order:
        if col != "Year" and col in new_annual.columns:
            new_annual[col] = pd.to_numeric(new_annual[col], errors="coerce").round(6)

    current_annual = page.production_annual.data[annual_order].copy()
    for col in annual_order:
        if col != "Year" and col in current_annual.columns:
            current_annual[col] = pd.to_numeric(current_annual[col], errors="coerce").round(6)

    if not new_annual.equals(current_annual):
        mark_user = user_edit_flag or user_driven_change or not page.production_annual.placeholder
        page.production_annual.set_data(new_annual, mark_user_input=mark_user)
        _update_table_editor_state(page.production_annual)
        _reset_table_widget(page.production_annual)
        updated = True

    if updated:
        _mark_inputs_dirty()


def _apply_growth_cascade(
    page: InputLandingPage,
    table: EditableTable,
    df: pd.DataFrame,
    row_idx: int,
    growth_percent: float,
) -> tuple[pd.DataFrame, Tuple[float, float] | None]:
    """Apply *growth_percent* (entered as percentage) to cascade production."""

    if df.empty or row_idx not in df.index:
        return df, None

    month_col = _get_month_column(df)
    if month_col is None or "Cassava ton" not in df.columns:
        return df, None

    month_values = pd.to_datetime(df[month_col].astype(str), errors="coerce").dt.to_period("M")
    if month_values.isna().all():
        return df, None

    try:
        base_period = month_values.iloc[row_idx]
    except Exception:  # pragma: no cover - defensive guard
        return df, None

    if pd.isna(base_period):
        return df, None

    rate_decimal = float(growth_percent)
    if not np.isfinite(rate_decimal):
        rate_decimal = 0.0
    if abs(rate_decimal) > 1.0:
        rate_decimal = rate_decimal / 100.0
    rate_decimal = float(np.clip(rate_decimal, -0.99, 10.0))

    working = df.copy()
    cassava_col = "Cassava ton"
    working[cassava_col] = pd.to_numeric(working[cassava_col], errors="coerce")
    growth_col = next((c for c in working.columns if "growth" in c.lower()), None)

    later_mask = month_values > base_period
    working.loc[later_mask, cassava_col] = np.nan

    if growth_col:
        working[growth_col] = pd.to_numeric(working[growth_col], errors="coerce")
        working.at[row_idx, growth_col] = rate_decimal
        working.loc[later_mask, growth_col] = np.nan

    table.set_data(working, mark_user_input=True)
    _update_table_editor_state(table)
    st.session_state[PRODUCTION_EDIT_FLAG] = True
    _mark_inputs_dirty()

    _auto_compound_production(page)

    refreshed = page.production_monthly.data.copy()
    derived: Tuple[float, float] | None = None
    refreshed_month_col = _get_month_column(refreshed)
    if refreshed_month_col and "Cassava ton" in refreshed.columns:
        month_str = base_period.strftime("%Y-%m")
        match = refreshed[refreshed[refreshed_month_col] == month_str]
        if not match.empty:
            cassava_val = pd.to_numeric(match["Cassava ton"], errors="coerce").iloc[0]
            if pd.notna(cassava_val):
                ethanol_val = float(cassava_val) * ETHANOL_LITRES_PER_TON
                feed_val = float(cassava_val) * ANIMAL_FEED_TON_PER_TON
                derived = (ethanol_val, feed_val)

    return refreshed, derived


def _update_staff_costs_from_positions(page: InputLandingPage) -> None:
    """Keep the monthly staff cost table aligned with position salaries."""

    if page.staff_positions.placeholder:
        return

    staff_df = page.staff_costs_monthly.data.copy()
    if staff_df.empty or "Department" not in staff_df.columns:
        return

    schedule = compute_staff_schedule(page.staff_positions.model_frame)
    summary = schedule.department_summary
    if summary.empty or "Average Monthly Salary" not in summary.columns:
        return

    avg_salary = summary.set_index("Department")["Average Monthly Salary"].to_dict()
    staff_df["Headcount"] = pd.to_numeric(staff_df.get("Headcount"), errors="coerce").fillna(0.0)

    updated_costs = []
    for _, row in staff_df.iterrows():
        dept = row.get("Department")
        salary = avg_salary.get(dept)
        if salary is None or not np.isfinite(salary):
            try:
                current_cost = float(row.get("Cost", 0.0))
            except (TypeError, ValueError):
                current_cost = 0.0
            updated_costs.append(current_cost)
        else:
            updated_costs.append(float(row.get("Headcount", 0.0)) * float(salary))

    staff_df["Cost"] = updated_costs
    page.staff_costs_monthly.set_data(staff_df, mark_user_input=True)


def _sync_working_capital_tables(page: InputLandingPage) -> None:
    """Refresh the Accounts Payable editor state after landing-page edits."""

    inv_table = page.inventory_payable
    if inv_table is None:
        return

    # With payables metrics now maintained exclusively in the Accounts Payable table,
    # we only need to refresh the editor cache so user edits remain visible. The
    # values themselves come directly from the landing-page entries.
    _update_table_editor_state(inv_table)


def _apply_dependent_updates(page: InputLandingPage) -> None:
    """Ensure derived landing-page tables stay synchronised with inputs."""

    _auto_compound_production(page)
    _update_staff_costs_from_positions(page)
    _sync_working_capital_tables(page)


def _key_assumptions_controls(table: EditableTable) -> None:
    """Expose frequently tweaked assumptions inside the main page."""

    st.subheader("Key Assumptions")
    original_df = table.data.copy()
    df = table.data.copy()
    updated = False
    slider_cfg = {
        "Corporate tax rate": dict(min_value=0.0, max_value=0.7, step=0.01),
        "Investor share capital": dict(min_value=0.0, max_value=1.0, step=0.01),
        "Owner share capital": dict(min_value=0.0, max_value=1.0, step=0.01),
        "Discount rate": dict(min_value=0.0, max_value=0.5, step=0.01),
    }

    for parameter, cfg in slider_cfg.items():
        if parameter in df["Parameter"].values:
            idx = df.index[df["Parameter"] == parameter][0]
            state_key = f"key_assumption_{parameter.replace(' ', '_').lower()}"
            value_key = f"{state_key}_value"
            min_value = float(cfg["min_value"])
            max_value = float(cfg["max_value"])
            step = float(cfg["step"])

            current_default = (
                float(df.at[idx, "Value"]) if pd.notna(df.at[idx, "Value"]) else min_value
            )
            if value_key not in st.session_state:
                st.session_state[value_key] = current_default

            original_value = current_default

            current_value = st.number_input(
                parameter,
                min_value=min_value,
                max_value=max_value,
                step=float(step),
                value=float(st.session_state[value_key]),
                key=f"{state_key}_input",
            )
            st.session_state[value_key] = float(current_value)
            df.at[idx, "Value"] = float(current_value)
            if not np.isclose(original_value, float(current_value)):
                updated = True
                _mark_inputs_dirty()
    table.set_data(df, mark_user_input=updated)
    if not df.equals(original_df):
        _update_table_editor_state(table)


def _numeric_step(value: float) -> float:
    """Return a sensible increment for Streamlit number inputs."""

    if value is None or pd.isna(value):
        return 0.01
    value = abs(float(value))
    if value == 0:
        return 0.01
    exponent = max(-2, int(np.floor(np.log10(value))) - 1)
    return round(10 ** exponent, 6)


def _display_production_metrics(derived_metrics: Tuple[float, float] | None) -> None:
    """Render helper metrics for production table edits."""

    if derived_metrics is None:
        st.info(
            "Enter a cassava tonnage to calculate the ethanol and animal feed outputs for this row."
        )
        return

    ethanol_val, feed_val = derived_metrics
    metric_cols = st.columns(2)
    metric_cols[0].metric(
        "Calculated Ethanol (litres)",
        f"{ethanol_val:,.0f}",
    )
    metric_cols[1].metric(
        "Calculated Animal Feed (ton)",
        f"{feed_val:,.3f}",
    )


def _production_metrics_from_row(df: pd.DataFrame, row_idx: int) -> Tuple[float, float] | None:
    if "Cassava ton" not in df.columns or row_idx not in df.index:
        return None

    cassava_val = pd.to_numeric(pd.Series([df.at[row_idx, "Cassava ton"]]), errors="coerce").iloc[0]
    if pd.isna(cassava_val):
        return None

    ethanol_val = float(cassava_val) * ETHANOL_LITRES_PER_TON
    feed_val = float(cassava_val) * ANIMAL_FEED_TON_PER_TON
    return ethanol_val, feed_val


def _sync_production_outputs(df: pd.DataFrame) -> pd.DataFrame:
    if "Cassava ton" not in df.columns:
        return df

    cassava_series = pd.to_numeric(df["Cassava ton"], errors="coerce")
    if "Ethanol litres" in df.columns:
        df.loc[:, "Ethanol litres"] = cassava_series * ETHANOL_LITRES_PER_TON
    if "Animal Feed ton" in df.columns:
        df.loc[:, "Animal Feed ton"] = cassava_series * ANIMAL_FEED_TON_PER_TON
    return df


def _row_editor_form(
    table: EditableTable,
    row_idx: int,
    projection: ProjectionHorizon,
    *,
    widget_prefix: str,
) -> Tuple[pd.DataFrame, bool, Tuple[float, float] | None]:
    """Shared routine to edit a single row within a landing-page table."""

    df = table.data.copy()
    if df.empty or row_idx not in df.index:
        return df, False, None

    original_row = df.loc[row_idx].copy()

    month_range = pd.period_range(
        f"{int(projection.start_year):04d}-01",
        f"{int(projection.end_year):04d}-12",
        freq="M",
    )
    month_options = [p.strftime("%Y-%m") for p in month_range]

    derived_columns = DERIVED_COLUMN_MAP.get(table.name, set())
    derived_metrics: Tuple[float, float] | None = None
    updated = False

    for column in table.columns:
        current_value = df.at[row_idx, column]
        widget_key = f"{widget_prefix}_{table.name}_{row_idx}_{column}".replace(" ", "_").lower()

        if column in derived_columns:
            if table.name == "Production Monthly":
                # Skip rendering derived production outputs so the focused editor
                # only exposes the cassava driver and growth inputs.
                continue

            display_value = ""
            if current_value is not None and not (isinstance(current_value, float) and pd.isna(current_value)):
                if isinstance(current_value, (int, float, np.floating, np.integer)):
                    display_value = f"{float(current_value):,.6f}".rstrip("0").rstrip(".")
                else:
                    display_value = str(current_value)
            st.text_input(
                column,
                value=display_value,
                key=widget_key,
                disabled=True,
                help="This value is calculated automatically from other inputs.",
            )
            continue

        column_label = column.lower().replace("_", " ")
        if re.search(r"\bmonth\b", column_label):
            current_str = (
                None
                if current_value is None or (isinstance(current_value, float) and pd.isna(current_value))
                else str(current_value)
            )

            option_list = ["Not set"] + month_options
            if current_str and current_str not in option_list:
                option_list.insert(1, current_str)

            default_index = 0
            if current_str and current_str in option_list:
                default_index = option_list.index(current_str)

            selection = st.selectbox(
                column,
                options=option_list,
                index=default_index,
                key=widget_key,
            )

            new_value = None if selection == "Not set" else selection
            df.at[row_idx, column] = new_value
            if current_str != new_value:
                updated = True
            continue

        category_key = (table.name, column)
        if category_key in CATEGORY_SELECT_OPTIONS:
            config = CATEGORY_SELECT_OPTIONS[category_key]
            options = list(config.get("options", []))
            allow_custom = bool(config.get("allow_custom", True))
            current_str = (
                ""
                if current_value is None or (isinstance(current_value, float) and pd.isna(current_value))
                else str(current_value)
            )
            if current_str and current_str not in options and allow_custom:
                options.append(current_str)

            if allow_custom:
                custom_label = "Custom value"
                if custom_label not in options:
                    options.append(custom_label)

            if options:
                if current_str and current_str in options:
                    default_index = options.index(current_str)
                else:
                    default_index = 0
            else:
                options = [""]
                default_index = 0

            selection = st.selectbox(
                column,
                options=options,
                index=default_index,
                key=widget_key,
            )

            if allow_custom and selection == "Custom value":
                custom_key = f"{widget_prefix}_{table.name}_{row_idx}_{column}_custom".replace(" ", "_").lower()
                custom_value = st.text_input(
                    f"Specify {column.lower()}",
                    value=current_str if current_str not in config.get("options", []) else "",
                    key=custom_key,
                )
                new_value = custom_value.strip() if custom_value.strip() else None
            else:
                new_value = selection
                if not allow_custom and new_value not in config.get("options", []):
                    new_value = config.get("options", [None])[0]

            df.at[row_idx, column] = new_value
            if (current_str or None) != (new_value or None):
                updated = True
            continue

        numeric_series = pd.to_numeric(df[column], errors="coerce")
        is_numeric = pd.api.types.is_numeric_dtype(df[column]) or numeric_series.notna().any()

        if is_numeric:
            if row_idx in numeric_series.index:
                base_value = numeric_series.loc[row_idx]
            else:
                base_value = numeric_series.iloc[0] if not numeric_series.empty else 0.0
            if pd.isna(base_value):
                base_value = 0.0
            original_value = float(base_value)
            step = float(_numeric_step(base_value))
            number_format = "%.0f" if step >= 1 else "%.4f"
            new_value = st.number_input(
                column,
                value=float(base_value),
                step=step,
                format=number_format,
                key=widget_key,
            )
            if pd.api.types.is_integer_dtype(df[column]):
                new_value = int(round(new_value))
            df.at[row_idx, column] = new_value
            if not np.isclose(original_value, float(new_value)):
                updated = True
        else:
            text_value = "" if current_value is None or pd.isna(current_value) else str(current_value)
            new_value = st.text_input(
                column,
                value=text_value,
                key=widget_key,
            )
            df.at[row_idx, column] = new_value
            if str(new_value) != str(text_value):
                updated = True

    if table.name == "Production Monthly" and "Cassava ton" in df.columns:
        cassava_raw = df.at[row_idx, "Cassava ton"]
        cassava_val = pd.to_numeric(pd.Series([cassava_raw]), errors="coerce").iloc[0]
        if pd.notna(cassava_val):
            ethanol_val = float(cassava_val) * ETHANOL_LITRES_PER_TON
            feed_val = float(cassava_val) * ANIMAL_FEED_TON_PER_TON
            if "Ethanol litres" in df.columns:
                df.at[row_idx, "Ethanol litres"] = ethanol_val
            if "Animal Feed ton" in df.columns:
                df.at[row_idx, "Animal Feed ton"] = feed_val
            derived_metrics = (ethanol_val, feed_val)

    if not df.loc[row_idx].equals(original_row):
        updated = True

    if updated and table.name == "Production Monthly":
        st.session_state[PRODUCTION_EDIT_FLAG] = True

    return df, updated, derived_metrics


def _modify_default_inputs(page: InputLandingPage) -> None:
    """Allow users to tweak any default input/figure via focused controls."""

    st.subheader("Modify Default Inputs & Figures")
    tables = page.tables()
    if not tables:
        st.info("No tables are available to edit.")
        return

    table_names = list(tables.keys())
    table_name = st.selectbox(
        "Select table",
        table_names,
        key="default_table_select",
    )
    table = tables[table_name]
    df = table.data.copy()

    if df.empty:
        st.info("The selected table has no rows to modify. Use the table editor below to add data.")
        return

    id_column = table.columns[0] if table.columns else None
    row_indices = list(df.index)

    def _format_row(idx: int) -> str:
        if id_column and id_column in table.data.columns:
            label = table.data.at[idx, id_column]
            if label is None or pd.isna(label) or str(label).strip() == "":
                label = f"Row {idx + 1}"
        else:
            label = f"Row {idx + 1}"

        if (
            table.name == "Direct Costs Monthly"
            and "Cost Category" in table.data.columns
        ):
            category = table.data.at[idx, "Cost Category"]
            if category and not pd.isna(category):
                return str(category)
            return f"Row {idx + 1}"

        return f"{idx + 1}. {label}"

    row_idx = st.selectbox(
        "Select row",
        row_indices,
        format_func=_format_row,
        key=f"default_row_select_{table_name}",
    )

    st.markdown("Adjust the values for the selected row:")

    df_updated, updated, derived_metrics = _row_editor_form(
        table,
        row_idx,
        page.projection,
        widget_prefix="default_edit",
    )

    increment_config = YEARLY_INCREMENT_CONFIG.get(table.name)
    if increment_config:
        available_columns = [col for col in increment_config["value_columns"] if col in df_updated.columns]
        if available_columns:
            st.markdown("---")
            st.markdown("**Yearly increment adjustments (%)**")
            description = increment_config.get("description")
            if description:
                st.caption(description)

            rates: Dict[str, float] = {}
            for column in available_columns:
                rate_input = st.number_input(
                    f"{column} annual change (%)",
                    value=0.0,
                    step=0.1,
                    format="%.2f",
                    key=f"yearly_increment_{table.name.replace(' ', '_').lower()}_{row_idx}_{column}",
                )
                rates[column] = float(rate_input) / 100.0

            if st.button(
                "Apply yearly increment",
                key=f"apply_yearly_increment_{table.name.replace(' ', '_').lower()}_{row_idx}",
            ):
                incremented = apply_yearly_increment(
                    df_updated,
                    row_idx,
                    date_column=increment_config["date_column"],
                    frequency=increment_config["frequency"],
                    value_columns=available_columns,
                    increments=rates,
                    match_columns=increment_config.get("match_columns", []),
                    horizon_end=_increment_horizon_value(increment_config, page.projection),
                )
                if table.name.startswith("Production"):
                    incremented = _sync_production_outputs(incremented)
                    derived_metrics = _production_metrics_from_row(incremented, row_idx)
                if not incremented.equals(df_updated):
                    df_updated = incremented
                    updated = True

    cascade_applied = False
    cascade_metrics: Tuple[float, float] | None = None

    if table.name == "Production Monthly":
        _display_production_metrics(derived_metrics)

        growth_col = next((c for c in df_updated.columns if "growth" in c.lower()), None)
        default_growth_pct = 0.0
        if growth_col and row_idx in df_updated.index:
            current_growth = pd.to_numeric(pd.Series([df_updated.at[row_idx, growth_col]]), errors="coerce").iloc[0]
            if pd.notna(current_growth):
                default_growth_pct = float(current_growth) * 100.0

        st.markdown("---")
        st.markdown("**Cascade growth across projection horizon**")
        st.caption(
            "Set a monthly growth percentage to apply from this month onward. The model will "
            "recalculate all subsequent production months automatically."
        )
        growth_input = st.number_input(
            "Monthly growth % to apply",
            value=float(default_growth_pct),
            step=0.1,
            format="%.2f",
            key=f"growth_pct_apply_{table.name.replace(' ', '_')}_{row_idx}",
        )
        if st.button(
            "Apply growth to remaining months",
            key=f"apply_growth_btn_{table.name.replace(' ', '_')}_{row_idx}",
        ):
            df_updated, cascade_metrics = _apply_growth_cascade(
                page,
                table,
                df_updated,
                row_idx,
                growth_input,
            )
            cascade_applied = True
            if cascade_metrics is not None:
                derived_metrics = cascade_metrics

    st.caption("Updates are applied immediately. Use the section tables below for bulk edits or row management.")
    if cascade_applied:
        df = table.data.copy()
        if not df.equals(df_updated):
            table.set_data(df, mark_user_input=True)
        _update_table_editor_state(table)
    elif updated:
        table.set_data(df_updated, mark_user_input=True)
        _mark_inputs_dirty()
        _update_table_editor_state(table)
    else:
        # Ensure the focused editor stays aligned even when only formatting changes occur.
        _update_table_editor_state(table)

    if table.placeholder and not table.data.empty:
        # Once a user opens the focused editor we treat the existing rows as
        # intentional input so downstream schedules stop ignoring them as
        # placeholders. This mirrors the manual "replace placeholder rows"
        # guidance we provide in the UI.
        table.set_data(table.data.copy(), mark_user_input=True)
        _update_table_editor_state(table)


def _apply_production_delta(
    page: InputLandingPage, table: EditableTable, row_idx: int, delta: float
) -> bool:
    """Adjust the cassava tonnage for *row_idx* by *delta* and refresh derived fields."""

    if table.name != "Production Monthly" or "Cassava ton" not in table.columns:
        return False

    df = table.data.copy()
    if df.empty or row_idx not in df.index:
        return False

    cassava_raw = pd.to_numeric(pd.Series([df.at[row_idx, "Cassava ton"]]), errors="coerce").iloc[0]
    current_value = float(cassava_raw) if pd.notna(cassava_raw) else 0.0
    new_value = max(0.0, current_value + float(delta))

    if not np.isfinite(new_value) or np.isclose(new_value, current_value, atol=1e-9):
        return False

    df.at[row_idx, "Cassava ton"] = new_value
    if "Ethanol litres" in df.columns:
        df.at[row_idx, "Ethanol litres"] = new_value * ETHANOL_LITRES_PER_TON
    if "Animal Feed ton" in df.columns:
        df.at[row_idx, "Animal Feed ton"] = new_value * ANIMAL_FEED_TON_PER_TON

    table.set_data(df, mark_user_input=True)
    st.session_state[PRODUCTION_EDIT_FLAG] = True
    _mark_inputs_dirty()
    _update_table_editor_state(table)
    return True


def _render_production_panel(page: InputLandingPage) -> None:
    """Expose production schedules outside the grouped tab layout."""

    st.subheader("Production Schedule")
    st.markdown(
        """
        - **Cascade growth across the projection horizon.**
        - Set a monthly growth percentage to apply from this month onward; the model will recalculate every subsequent production month automatically.
        - Updates are applied immediately—use the section tables below for bulk edits or row management.
        - The production schedule you see here is derived from your changes, not the seeded defaults.
        - Adjust cassava tonnage with the ± controls or insert dated changes on the right. Monthly figures automatically cascade to the annual production summary.
        """
    )

    monthly_table = page.production_monthly
    annual_table = page.production_annual

    monthly_df = monthly_table.data.copy()

    change_applied = False
    if not monthly_df.empty:
        month_col = _get_month_column(monthly_df) or monthly_table.columns[0]

        def _format_row(idx: int) -> str:
            if month_col in monthly_df.columns:
                value = monthly_df.at[idx, month_col]
                if pd.notna(value):
                    return str(value)
            return f"Row {idx + 1}"

        row_idx = st.selectbox(
            "Month to adjust",
            list(monthly_df.index),
            format_func=_format_row,
            key="production_month_select",
        )

        step_value = st.number_input(
            "Adjustment step (tons)",
            min_value=0.0,
            value=100.0,
            step=10.0,
            key="production_step_value",
        )

        growth_col = next((c for c in monthly_df.columns if "growth" in c.lower()), None)
        default_growth_pct = 0.0
        if growth_col is not None and row_idx in monthly_df.index:
            current_growth = pd.to_numeric(
                pd.Series([monthly_df.at[row_idx, growth_col]]), errors="coerce"
            ).iloc[0]
            if pd.notna(current_growth):
                default_growth_pct = float(current_growth) * 100.0

        controls = st.columns([1, 1, 2, 2])
        minus_pressed = controls[0].button("−", key=f"production_minus_{row_idx}")
        plus_pressed = controls[1].button("+", key=f"production_plus_{row_idx}")

        with controls[2]:
            st.markdown("**Monthly growth % to apply**")
            growth_value = st.number_input(
                "Monthly growth % to apply",
                value=float(default_growth_pct),
                step=0.1,
                format="%.2f",
                key=f"production_growth_pct_{row_idx}",
            )
        with controls[3]:
            if st.button(
                "Cascade growth",
                key=f"production_cascade_growth_{row_idx}",
                help="Apply the selected growth rate from this month onward.",
            ):
                _apply_growth_cascade(
                    page,
                    monthly_table,
                    monthly_table.data.copy(),
                    row_idx,
                    growth_value,
                )
                change_applied = True

        st.markdown("**Change effective month**")
        st.caption(
            "Select the month where the new production plan should begin, click **Add change**, and edit the inserted row. "
            "The model will cascade cassava tonnage (and the derived ethanol and animal feed outputs) to all later months automatically."
        )
        change_controls_applied = _render_change_controls(page, monthly_table)
        change_applied = change_applied or change_controls_applied

        if step_value > 0:
            if minus_pressed:
                _apply_production_delta(page, monthly_table, row_idx, -step_value)
            if plus_pressed:
                _apply_production_delta(page, monthly_table, row_idx, step_value)
        elif minus_pressed or plus_pressed:
            st.warning("Set a step above zero to apply the adjustment.")

    else:
        st.info(
            "The production schedule is empty. Use the bulk change controls to insert the first production month."
        )
        st.markdown("**Change effective month**")
        st.caption(
            "Select the month where the new production plan should begin, click **Add change**, and edit the newly inserted row. "
            "The model will cascade cassava tonnage (and the derived ethanol and animal feed outputs) to all later months automatically."
        )
        change_applied = _render_change_controls(page, monthly_table)

    if change_applied and monthly_table.name == "Production Monthly":
        st.session_state[PRODUCTION_EDIT_FLAG] = True

    _render_table(
        page,
        monthly_table,
        "Production",
        expanded=True,
        allow_change_controls=False,
    )

    _render_table(
        page,
        annual_table,
        "Production",
        expanded=False,
        allow_change_controls=False,
    )


def _editable_tables(page: InputLandingPage) -> None:
    """Render editable data tables grouped by the landing-page sections."""

    categories = page.grouped_tables()
    filtered = OrderedDict((k, v) for k, v in categories.items() if k != "Production")
    if not filtered:
        return

    tabs = st.tabs(list(filtered.keys()))

    for tab, (section, tables) in zip(tabs, filtered.items()):
        with tab:
            for table in tables:
                expanded = section in {"Global", "Capex", "Financial"}
                _render_table(page, table, section, expanded=expanded)
                if table is page.initial_investment:
                    st.metric("Total Initial Investment", _format_currency(page.total_initial_investment))


def _render_table(
    page: InputLandingPage,
    table: EditableTable,
    section: str,
    *,
    expanded: bool = False,
    allow_change_controls: bool = True,
) -> None:
    """Show a Streamlit data editor for a specific table."""

    safe_key = table.name.replace(' ', '_').lower()
    widget_key = _table_widget_key(table)
    cache_key = _table_editor_state_key(table)
    with st.expander(table.name, expanded=expanded):
        original_data = table.data.copy()
        data_changed = False
        controls = st.columns(2)
        if controls[0].button("➕ Add row", key=f"add_{safe_key}"):
            table.add_row({column: None for column in table.columns})
            data_changed = True

        if not table.data.empty:
            row_options = list(table.data.index)
            remove_index = controls[1].selectbox(
                "Row to remove",
                options=row_options,
                key=f"remove_select_{safe_key}",
                label_visibility="collapsed",
                format_func=lambda i: f"Row {i + 1}",
            )
            if controls[1].button("➖ Remove selected", key=f"remove_{safe_key}"):
                idx = int(remove_index)
                table.remove_row(idx)
                data_changed = True
        else:
            controls[1].markdown("&nbsp;")

        change_applied = False
        if allow_change_controls:
            change_applied = _render_change_controls(page, table)
            if change_applied:
                data_changed = True

        derived_columns = DERIVED_COLUMN_MAP.get(table.name, set())
        column_config = {}
        for col in derived_columns:
            if col in table.columns:
                column_config[col] = st.column_config.NumberColumn(
                    label=col,
                    disabled=True,
                    help="Calculated automatically from other inputs.",
                )

        for (table_name, column), config in CATEGORY_SELECT_OPTIONS.items():
            if (
                table_name == table.name
                and column in table.columns
                and not config.get("allow_custom", True)
            ):
                column_config[column] = st.column_config.SelectboxColumn(
                    label=column,
                    options=list(config.get("options", [])),
                )

        if table.name == "Production Monthly":
            st.caption(
                "Edit **Cassava ton** (and optional Growth %) for any month. The model will automatically recompute the matching "
                "ethanol litres and animal-feed tonnage and roll the results into the annual production table."
            )
        elif table.name == "Production Annual":
            st.caption(
                "These annual totals are derived from the monthly production schedule. Update the monthly table to change the "
                "values shown here."
            )

        if cache_key not in st.session_state:
            st.session_state[cache_key] = table.data.copy()

        edited = st.data_editor(
            table.data,
            num_rows="dynamic",
            use_container_width=True,
            key=widget_key,
            column_config=column_config or None,
        )
        if isinstance(edited, pd.DataFrame):
            new_data = edited[table.columns].copy()
        else:  # pragma: no cover - safety for older Streamlit returning list of dicts
            new_data = pd.DataFrame(edited, columns=table.columns)

        if not new_data.equals(table.data):
            data_changed = True

        table.set_data(new_data, mark_user_input=data_changed)
        if (data_changed or change_applied) and table.name == "Production Monthly":
            st.session_state[PRODUCTION_EDIT_FLAG] = True
        _update_table_editor_state(table)

        if data_changed or not table.data.equals(original_data):
            _mark_inputs_dirty()

def _annualise(rate: float | None) -> float | None:
    if rate is None or pd.isna(rate):
        return None
    return (1 + rate) ** 12 - 1


def _format_rate(rate: float | None) -> str:
    if rate is None or pd.isna(rate):
        return "n/a"
    return f"{rate * 100:,.2f}%"


def _format_percent(value: float | None) -> str:
    if value is None or pd.isna(value):
        return "n/a"
    return f"{value * 100:,.1f}%"


def _format_currency(value: float | None) -> str:
    if value is None or pd.isna(value):
        return "n/a"
    return f"${value:,.0f}"


def _reset_period_index(df: pd.DataFrame, label: str) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame(columns=[label])
    result = df.copy().reset_index()
    if result.columns[0] != label:
        result = result.rename(columns={result.columns[0]: label})
    if np.issubdtype(result[label].dtype, np.datetime64):
        result[label] = pd.to_datetime(result[label]).dt.to_period("M").astype(str)
    return result

def _render_key_metrics(model: CassavaBioethanolModel, results: Dict[str, object]) -> None:
    metrics = results["metrics"]
    revenue = results["revenue"]
    production = results["production"]
    costs = results["costs"]
    expenses_summary: ExpenseSummary | None = results.get("expenses") if isinstance(results.get("expenses"), ExpenseSummary) else None
    financials = results["financials"]
    loan_schedule = results["loan_schedule"]
    depreciation = results["depreciation"]

    st.subheader("Overview")
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Project NPV", _format_currency(metrics.get("Project NPV")))
    col2.metric("Project IRR (annual)", _format_rate(_annualise(metrics.get("Project IRR"))))
    col3.metric("Equity IRR (annual)", _format_rate(_annualise(metrics.get("Equity IRR"))))
    payback_years = metrics.get("Payback Period (years)")
    col4.metric("Payback Period (years)", f"{payback_years:,.1f}" if payback_years and not pd.isna(payback_years) else "n/a")

    st.markdown("### Assumptions Snapshot")
    assumption_snapshot = pd.DataFrame(
        {
            "Assumption": [
                "Corporate Tax Rate",
                "Investor Share",
                "Owner Share",
                "Discount Rate",
                "Terminal Growth Rate",
                "Capital Gains Tax Rate",
                "Total Initial Investment",
                "Initial Equity Investment",
                "Initial Loan Funding",
            ],
            "Value": [
                _format_rate(metrics.get("Corporate Tax Rate")),
                _format_percent(metrics.get("Investor Share")),
                _format_percent(metrics.get("Owner Share")),
                _format_rate(metrics.get("Discount Rate")),
                _format_rate(metrics.get("Terminal Growth Rate")),
                _format_rate(metrics.get("Capital Gains Tax Rate")),
                _format_currency(metrics.get("Total Initial Investment")),
                _format_currency(metrics.get("Initial Equity Investment")),
                _format_currency(metrics.get("Initial Loan Funding")),
            ],
        }
    )
    st.dataframe(assumption_snapshot, use_container_width=True, hide_index=True)

    st.markdown("### Detailed Working Capital (Annual)")
    working_capital_output = results.get("working_capital")
    if isinstance(working_capital_output, WorkingCapitalOutput):
        annual_wc = working_capital_output.annual.copy()
    else:
        annual_wc = pd.DataFrame()
    if isinstance(annual_wc, pd.DataFrame) and not annual_wc.empty:
        annual_display = annual_wc.copy()
        annual_display.index.name = "Year"
        st.dataframe(annual_display.reset_index(), use_container_width=True)
    else:
        st.info("Working capital schedule is not available for the selected horizon.")

    st.markdown("### Latest Drivers")
    drivers = pd.DataFrame(
        {
            "Metric": [
                "Final Month Revenue",
                "Final Month EBITDA",
                "Final Month Equity Cash Flow",
                "Cumulative FCF",
                "Cumulative Equity CF",
            ],
            "Value": [
                _format_currency(metrics.get("Final Month Revenue")),
                _format_currency(metrics.get("Final Month EBITDA")),
                _format_currency(metrics.get("Final Month Equity CF")),
                _format_currency(metrics.get("Cumulative FCF")),
                _format_currency(metrics.get("Cumulative Equity CF")),
            ],
        }
    )
    st.dataframe(drivers, use_container_width=True, hide_index=True)

    st.markdown("### Lender Coverage Metrics")
    lender_df = pd.DataFrame(
        {
            "Metric": [
                "DSCR (min)",
                "DSCR (avg)",
                "Debt Service Coverage Breach Months",
                "LLCR",
                "PLCR",
                "Outstanding Debt (opening)",
                "PV CFADS",
                "PV Terminal Value",
            ],
            "Value": [
                _format_rate(metrics.get("DSCR (min)")),
                _format_rate(metrics.get("DSCR (avg)")),
                f"{metrics.get('Debt Service Coverage Breach Months', 0):,.0f}",
                _format_rate(metrics.get("LLCR")),
                _format_rate(metrics.get("PLCR")),
                _format_currency(metrics.get("Outstanding Debt (opening)")),
                _format_currency(metrics.get("PV CFADS")),
                _format_currency(metrics.get("PV Terminal Value")),
            ],
        }
    )
    st.dataframe(lender_df, use_container_width=True, hide_index=True)

    st.markdown("### Covenant Heatmap (DSCR)")
    dscr_base = pd.to_numeric(financials.cashflow_monthly.get("Operating Cash Flow"), errors="coerce").fillna(0.0)
    dscr_debt = pd.to_numeric(financials.cashflow_monthly.get("Debt Service"), errors="coerce").replace(0.0, np.nan)
    dscr_series = (dscr_base / dscr_debt).replace([np.inf, -np.inf], np.nan)
    if not dscr_series.dropna().empty:
        dscr_df = dscr_series.to_frame("DSCR")
        dscr_df["Month"] = dscr_df.index.to_period("M").astype(str)
        dscr_df["Breach (<1.0x)"] = dscr_df["DSCR"] < 1.0
        heat = px.imshow(
            np.array([dscr_df["DSCR"].fillna(0.0).values]),
            aspect="auto",
            color_continuous_scale="RdYlGn",
            labels={"x": "Month", "y": "Metric", "color": "DSCR"},
        )
        heat.update_xaxes(tickmode="array", tickvals=list(range(len(dscr_df))), ticktext=dscr_df["Month"].tolist())
        heat.update_yaxes(tickmode="array", tickvals=[0], ticktext=["DSCR"])
        st.plotly_chart(heat, use_container_width=True)
        st.dataframe(dscr_df[["Month", "DSCR", "Breach (<1.0x)"]], use_container_width=True, hide_index=True)
    else:
        st.info("No DSCR series available for covenant heatmap.")

    st.markdown("### Annual Operations & Production")
    production_annual = production.annual.copy()
    if not production_annual.empty:
        chart_data = production_annual.select_dtypes(include=[np.number])
        if not chart_data.empty:
            st.bar_chart(chart_data)
        else:
            st.info("No numeric production data available for charting.")
    else:
        st.info("No production data available for the selected horizon.")

    summary_cols = [col for col in ["Revenue", "EBITDA", "Net Income"] if col in financials.income_annual.columns]
    if summary_cols:
        annual_summary = financials.income_annual[summary_cols].copy()
        annual_summary.index.name = "Year"
        st.dataframe(annual_summary.reset_index(), use_container_width=True)

    st.markdown("### Cash Flow & Returns")
    cash_columns = [
        c
        for c in [
            "Operating Cash Flow",
            "Investing Cash Flow",
            "Financing Cash Flow",
            "Free Cash Flow",
            "Equity Cash Flow",
        ]
        if c in financials.cashflow_monthly.columns
    ]
    if cash_columns:
        st.line_chart(financials.cashflow_monthly[cash_columns])
        cumulative_df = financials.cashflow_monthly[cash_columns].cumsum()
        cumulative_df.columns = [f"Cumulative {col}" for col in cumulative_df.columns]
        st.line_chart(cumulative_df)

    st.markdown("### Revenue Mix")
    revenue_annual = revenue.annual.copy()
    if not revenue_annual.empty:
        if "Total Revenue" in revenue_annual:
            mix_df = revenue_annual.drop(columns=["Total Revenue"])
        else:
            mix_df = revenue_annual
        if not mix_df.empty:
            st.bar_chart(mix_df)
        st.dataframe(revenue_annual.reset_index().rename(columns={"index": "Year"}), use_container_width=True)
    else:
        st.info("Revenue inputs are empty for the current projection.")

    st.markdown("### Operating Cost Breakdown")
    if isinstance(expenses_summary, ExpenseSummary) and not expenses_summary.monthly.empty:
        st.area_chart(expenses_summary.monthly)
    else:
        cost_monthly = pd.DataFrame(
            {
                name: output.monthly.sum(axis=1)
                for name, output in costs.items()
                if output and not output.monthly.empty
            }
        )
        if not cost_monthly.empty:
            st.area_chart(cost_monthly)

    if isinstance(expenses_summary, ExpenseSummary) and not expenses_summary.annual.empty:
        annual_expense = expenses_summary.annual.copy()
        annual_expense.index.name = "Year"
        st.dataframe(annual_expense.reset_index(), use_container_width=True)
    else:
        cost_annual = pd.DataFrame(
            {
                name: output.annual.sum(axis=1)
                for name, output in costs.items()
                if output and not output.annual.empty
            }
        )
        if not cost_annual.empty:
            cost_annual.index.name = "Year"
            st.dataframe(cost_annual.reset_index(), use_container_width=True)

    st.markdown("### Capital Expenditure & Debt")
    capex_df = model.input_page.initial_investment.model_frame
    if not capex_df.empty:
        st.bar_chart(capex_df.set_index("Item")["Cost"])
        st.dataframe(capex_df, use_container_width=True)
    debt_chart = loan_schedule.schedule.pivot_table(index="Month", values="Closing Balance", aggfunc="sum")
    if not debt_chart.empty:
        st.line_chart(debt_chart)

    st.markdown("### Fixed Asset Summary")
    st.dataframe(depreciation.summary, use_container_width=True)

    st.markdown("### Debt Schedule Chart")
    debt_payments = loan_schedule.schedule.pivot_table(index="Month", values="Payment", aggfunc="sum")
    if not debt_payments.empty:
        st.area_chart(debt_payments)

    st.markdown("### Yearly Loan Amortisation")
    yearly_amortisation = getattr(loan_schedule, "annual", pd.DataFrame())
    if isinstance(yearly_amortisation, pd.DataFrame) and not yearly_amortisation.empty:
        amort_display = yearly_amortisation.copy()
        rate_column = "Interest Rate"
        currency_columns = [
            "Yearly Remaining Balance",
            "Monthly Interest (Balance × Rate / 12)",
            "Interest Paid",
            "Principal Paid",
            "Total Payment",
            "Year-End Balance",
        ]
        if rate_column in amort_display.columns:
            amort_display[rate_column] = amort_display[rate_column].apply(_format_rate)
        for col in currency_columns:
            if col in amort_display.columns:
                amort_display[col] = amort_display[col].apply(_format_currency)
        st.dataframe(amort_display, use_container_width=True)
    else:
        st.info("No loan amortisation data available for the current projection.")

    st.markdown("### Break-even Analysis")
    break_even_df = results.get("break_even")
    if isinstance(break_even_df, pd.DataFrame) and not break_even_df.empty:
        st.dataframe(_reset_period_index(break_even_df, "Month"), use_container_width=True)

        break_even_monthly = break_even_df.copy()
        if isinstance(break_even_monthly.index, pd.PeriodIndex):
            break_even_monthly.index = break_even_monthly.index.to_timestamp()
        elif not isinstance(break_even_monthly.index, pd.DatetimeIndex):
            converted_index = pd.to_datetime(break_even_monthly.index, errors="coerce")
            valid_mask = converted_index.notna()
            break_even_monthly = break_even_monthly.loc[valid_mask]
            break_even_monthly.index = converted_index[valid_mask]

        numeric_columns = [
            column
            for column in ("Monthly Margin", "Cumulative Margin")
            if column in break_even_monthly.columns
        ]

        if numeric_columns:
            monthly_chart = break_even_monthly[numeric_columns]
            if not monthly_chart.empty:
                st.markdown("#### Monthly Break-even Trend")
                st.line_chart(monthly_chart)

                aggregation: Dict[str, str] = {}
                if "Monthly Margin" in monthly_chart.columns:
                    aggregation["Monthly Margin"] = "sum"
                if "Cumulative Margin" in monthly_chart.columns:
                    aggregation["Cumulative Margin"] = "last"
                annual_break_even = (
                    monthly_chart.resample("Y").agg(aggregation).dropna(how="all") if aggregation else pd.DataFrame()
                )
                if not annual_break_even.empty:
                    annual_break_even.index = annual_break_even.index.to_period("Y").astype(str)
                    st.markdown("#### Annual Break-even Trend")
                    st.line_chart(annual_break_even)

def _render_financial_performance(results: Dict[str, object]) -> None:
    financials = results["financials"]
    expenses_summary: ExpenseSummary | None = (
        results.get("expenses") if isinstance(results.get("expenses"), ExpenseSummary) else None
    )
    costs = results["costs"]

    st.subheader("Monthly Financial Performance")
    st.dataframe(_reset_period_index(financials.income_monthly, "Month"), use_container_width=True)

    st.subheader("Annual Financial Performance")
    annual_income = financials.income_annual.copy()
    annual_income.index.name = "Year"
    st.dataframe(annual_income.reset_index(), use_container_width=True)

    expense_columns = ["COGS", "Staff Costs", "Other Opex"]

    def _expense_breakdown(source: pd.DataFrame, index_label: str) -> pd.DataFrame:
        subset_cols = [c for c in expense_columns if c in source.columns]
        frame = source.reindex(columns=subset_cols).copy()
        frame = frame.apply(pd.to_numeric, errors="coerce").fillna(0.0)
        if not frame.empty:
            frame["Total"] = frame.sum(axis=1)
        if index_label == "Year":
            frame.index.name = "Year"
        return frame

    if isinstance(expenses_summary, ExpenseSummary) and not expenses_summary.monthly.empty:
        monthly_expense_breakdown = _expense_breakdown(expenses_summary.monthly, "Month")
    else:
        monthly_expense_breakdown = _expense_breakdown(financials.income_monthly, "Month")
        if monthly_expense_breakdown.empty:
            monthly_expense_breakdown = pd.DataFrame(0.0, index=financials.income_monthly.index, columns=expense_columns + ["Total"])
    st.subheader("Expense Breakdown (Monthly)")
    st.dataframe(_reset_period_index(monthly_expense_breakdown, "Month"), use_container_width=True)

    if isinstance(expenses_summary, ExpenseSummary) and not expenses_summary.annual.empty:
        annual_expense_breakdown = _expense_breakdown(expenses_summary.annual, "Year")
    else:
        annual_expense_breakdown = _expense_breakdown(financials.income_annual, "Year")
        if annual_expense_breakdown.empty:
            annual_expense_breakdown = pd.DataFrame(0.0, index=financials.income_annual.index, columns=expense_columns + ["Total"])
            annual_expense_breakdown.index.name = "Year"
    st.subheader("Expense Breakdown (Annual)")
    st.dataframe(_reset_period_index(annual_expense_breakdown, "Year"), use_container_width=True)

    income_ratios_monthly = getattr(financials, "income_ratios_monthly", pd.DataFrame())
    if isinstance(income_ratios_monthly, pd.DataFrame) and not income_ratios_monthly.empty:
        ratio_monthly = _reset_period_index(income_ratios_monthly, "Month")
        if "Month" in ratio_monthly.columns:
            try:
                ratio_monthly["Month"] = pd.to_datetime(ratio_monthly["Month"]).dt.to_period("M").astype(str)
            except Exception:
                ratio_monthly["Month"] = ratio_monthly["Month"].astype(str)
        st.subheader("Income Statement Ratios (Monthly)")
        st.dataframe(ratio_monthly, use_container_width=True)

    income_ratios_annual = getattr(financials, "income_ratios_annual", pd.DataFrame())
    if isinstance(income_ratios_annual, pd.DataFrame) and not income_ratios_annual.empty:
        ratio_annual = _reset_period_index(income_ratios_annual, "Year")
        st.subheader("Income Statement Ratios (Annual)")
        st.dataframe(ratio_annual, use_container_width=True)

    st.subheader("Total Expense Schedule")
    if isinstance(expenses_summary, ExpenseSummary):
        if not expenses_summary.monthly.empty:
            st.dataframe(
                _reset_period_index(expenses_summary.monthly, "Month"),
                use_container_width=True,
            )
        if not expenses_summary.annual.empty:
            annual_expense = expenses_summary.annual.copy()
            annual_expense.index.name = "Year"
            st.dataframe(annual_expense.reset_index(), use_container_width=True)
    else:
        monthly_expense = pd.DataFrame(
            {
                name: output.monthly.sum(axis=1)
                for name, output in costs.items()
                if output and not output.monthly.empty
            }
        )
        if not monthly_expense.empty:
            st.dataframe(_reset_period_index(monthly_expense, "Month"), use_container_width=True)
        annual_expense = pd.DataFrame(
            {
                name: output.annual.sum(axis=1)
                for name, output in costs.items()
                if output and not output.annual.empty
            }
        )
        if not annual_expense.empty:
            annual_expense.index.name = "Year"
            st.dataframe(annual_expense.reset_index(), use_container_width=True)

    staff_schedule = results.get("staff_schedule")
    if staff_schedule is not None:
        positions_df = getattr(staff_schedule, "positions", pd.DataFrame())
        summary_df = getattr(staff_schedule, "department_summary", pd.DataFrame())
        if isinstance(positions_df, pd.DataFrame) and not positions_df.empty:
            st.subheader("Staff Position Schedule")
            st.dataframe(positions_df, use_container_width=True)
        if isinstance(summary_df, pd.DataFrame) and not summary_df.empty:
            st.subheader("Staff Cost by Department")
            st.dataframe(summary_df, use_container_width=True)

def _render_financial_position(results: Dict[str, object]) -> None:
    financials = results["financials"]

    st.subheader("Monthly Statement of Financial Position")
    st.dataframe(_reset_period_index(financials.balance_monthly, "Month"), use_container_width=True)

    st.subheader("Annual Statement of Financial Position")
    balance_annual = financials.balance_annual.copy()
    balance_annual.index.name = "Year"
    st.dataframe(balance_annual.reset_index(), use_container_width=True)

    balance_ratios_monthly = getattr(financials, "balance_ratios_monthly", pd.DataFrame())
    if isinstance(balance_ratios_monthly, pd.DataFrame) and not balance_ratios_monthly.empty:
        ratio_monthly = _reset_period_index(balance_ratios_monthly, "Month")
        if "Month" in ratio_monthly.columns:
            try:
                ratio_monthly["Month"] = pd.to_datetime(ratio_monthly["Month"]).dt.to_period("M").astype(str)
            except Exception:
                ratio_monthly["Month"] = ratio_monthly["Month"].astype(str)
        st.subheader("Statement of Financial Position Ratios (Monthly)")
        st.dataframe(ratio_monthly, use_container_width=True)

    balance_ratios_annual = getattr(financials, "balance_ratios_annual", pd.DataFrame())
    if isinstance(balance_ratios_annual, pd.DataFrame) and not balance_ratios_annual.empty:
        ratio_annual = _reset_period_index(balance_ratios_annual, "Year")
        st.subheader("Statement of Financial Position Ratios (Annual)")
        st.dataframe(ratio_annual, use_container_width=True)

def _render_cash_flow_page(results: Dict[str, object]) -> None:
    financials = results["financials"]

    st.subheader("Monthly Cash Flow Statement")
    cash_monthly = financials.cashflow_monthly
    st.dataframe(_reset_period_index(cash_monthly, "Month"), use_container_width=True)

    st.subheader("Annual Cash Flow Statement")
    cash_annual = financials.cashflow_annual.copy()
    cash_annual.index.name = "Year"
    st.dataframe(cash_annual.reset_index(), use_container_width=True)

    st.subheader("Cash Flow and Returns Charts")
    cash_columns = [
        c
        for c in [
            "Operating Cash Flow",
            "Investing Cash Flow",
            "Financing Cash Flow",
            "Free Cash Flow",
            "Equity Cash Flow",
        ]
        if c in cash_monthly.columns
    ]
    if cash_columns:
        st.line_chart(cash_monthly[cash_columns])

    st.subheader("Cumulative Equity Cash Flow")
    if "Equity Cash Flow" in cash_monthly.columns:
        cumulative_equity = cash_monthly[["Equity Cash Flow"]].cumsum()
        cumulative_equity.columns = ["Cumulative Equity Cash Flow"]
        st.line_chart(cumulative_equity)
        cumulative_df = cumulative_equity.reset_index().rename(columns={cumulative_equity.index.name or "index": "Month"})
        st.dataframe(cumulative_df, use_container_width=True)

def _render_sensitivity_page(model: CassavaBioethanolModel, results: Dict[str, object]) -> None:
    st.subheader("Sensitivity Analysis Configuration")
    config_df = pd.DataFrame([s.__dict__ for s in DEFAULT_SENSITIVITY_SCENARIOS]) if DEFAULT_SENSITIVITY_SCENARIOS else pd.DataFrame(columns=["name", "parameter", "delta"])
    st.dataframe(config_df.rename(columns={"name": "Scenario", "parameter": "Parameter", "delta": "Delta"}), use_container_width=True, hide_index=True)

    base_page = copy.deepcopy(results.get("input_page_snapshot", model.input_page))

    def _scenario_model() -> CassavaBioethanolModel:
        clone = CassavaBioethanolModel(copy.deepcopy(base_page))
        clone.scenario = model.scenario
        return clone

    analysis_model = _scenario_model()
    cache: Dict[str, Dict[str, object]] = st.session_state.setdefault(SENSITIVITY_CACHE_KEY, {})
    cached_entry = cache.get(model.scenario)
    run_requested = st.button(
        "Run Sensitivity Analysis",
        key=f"run_sensitivity_{model.scenario.lower()}",
    )

    if run_requested or not cached_entry or cached_entry.get("version") != _current_model_version():
        if DEFAULT_SENSITIVITY_SCENARIOS:
            with st.spinner("Running sensitivity cases..."):
                sensitivity_results = run_sensitivity(analysis_model, DEFAULT_SENSITIVITY_SCENARIOS)
        else:
            sensitivity_results = pd.DataFrame(
                columns=["Scenario", "Parameter", "Delta", "Project NPV", "Change vs Base"]
            )
        cache[model.scenario] = {
            "version": _current_model_version(),
            "results": sensitivity_results,
        }
        st.session_state[SENSITIVITY_CACHE_KEY] = cache
        cached_entry = cache[model.scenario]
    elif cached_entry is None:
        sensitivity_results = pd.DataFrame(
            columns=["Scenario", "Parameter", "Delta", "Project NPV", "Change vs Base"]
        )
    else:
        sensitivity_results = cached_entry.get("results", pd.DataFrame())

    st.subheader("Simulation Results")
    if sensitivity_results.empty:
        st.info("Click 'Run Sensitivity Analysis' to generate comparison results.")
    else:
        st.dataframe(sensitivity_results, use_container_width=True)

    st.subheader("Tornado Drivers")
    tornado_model = _scenario_model()
    tornado_df = tornado_chart_inputs(tornado_model, TORNADO_DRIVERS, scale=0.1)
    st.dataframe(tornado_df, use_container_width=True)

def _render_scenario_page(model: CassavaBioethanolModel, results: Dict[str, object]) -> None:
    base_page = copy.deepcopy(results.get("input_page_snapshot", model.input_page))
    parameter_catalog = scenario_parameter_catalog(base_page)

    st.subheader("Scenario Parameter Library")
    if parameter_catalog.empty:
        st.info("No scenario-compatible parameters are available for the current inputs.")
    else:
        catalog_display = parameter_catalog.copy()
        st.dataframe(catalog_display, use_container_width=True, hide_index=True)

    options = parameter_catalog["Parameter"].tolist()
    scenario_definitions = st.session_state.setdefault(SCENARIO_DEFINITIONS_KEY, [])

    st.subheader("Credit Committee Scenario Pack")
    st.caption(
        "Load pre-baked committee scenarios with correlated stresses (lower ethanol price proxy, "
        "higher feedstock costs, and ramp-delay proxy)."
    )
    if st.button("Load Base / Downside / Severe Downside / Upside", key="load_credit_committee_pack"):
        committee_configs = credit_committee_scenario_configs(base_page)
        scenario_definitions = [
            {
                "name": cfg.name,
                "overrides": cfg.overrides,
                "deltas": {k: None for k in cfg.overrides},
            }
            for cfg in committee_configs
        ]
        st.session_state[SCENARIO_DEFINITIONS_KEY] = scenario_definitions
        st.success("Credit committee scenario pack loaded.")
        _trigger_rerun()

    selected_defaults = st.session_state.setdefault(SCENARIO_SELECTION_STATE_KEY, [])
    selected_parameters = st.multiselect(
        "Select parameters to configure",
        options=options,
        default=[value for value in selected_defaults if value in options],
    )
    st.session_state[SCENARIO_SELECTION_STATE_KEY] = selected_parameters

    builder_values = st.session_state.setdefault(SCENARIO_VALUE_STATE_KEY, {})
    catalog_lookup = {row["Parameter"]: row for row in parameter_catalog.to_dict("records")}

    st.session_state.setdefault(SCENARIO_NAME_STATE_KEY, "")
    if st.session_state.pop(SCENARIO_CLEAR_NAME_FLAG, False):
        st.session_state[SCENARIO_NAME_STATE_KEY] = ""

    if selected_parameters:
        st.subheader("Scenario Builder")

    for parameter in selected_parameters:
        info = catalog_lookup.get(parameter, {})
        base_value = float(info.get("Base Value", np.nan)) if info else np.nan
        units = info.get("Units", "")
        source = info.get("Source", "")
        base_label = f"{base_value:,.2f}" if np.isfinite(base_value) else "n/a"
        expander_label = f"{parameter} — Base {base_label}{f' {units}' if units else ''}"
        with st.expander(expander_label, expanded=False):
            if source:
                st.caption(f"Source: {source}")
            default_mode = "percent" if np.isfinite(base_value) and not np.isclose(base_value, 0.0) else "absolute"
            state = builder_values.get(parameter, {"mode": default_mode})
            mode = st.selectbox(
                "Adjustment mode",
                ["Percent", "Absolute"],
                index=0 if state.get("mode", default_mode) == "percent" else 1,
                key=f"scenario_mode_{parameter}",
            )
            state["mode"] = "percent" if mode == "Percent" else "absolute"
            if state["mode"] == "percent" and np.isfinite(base_value) and not np.isclose(base_value, 0.0):
                percent_default = float(state.get("percent", 0.0))
                percent_value = st.number_input(
                    "Percent change (%)",
                    key=f"scenario_percent_{parameter}",
                    value=percent_default,
                    step=0.5,
                    format="%.2f",
                )
                state["percent"] = percent_value
            else:
                if not np.isfinite(base_value):
                    st.info("Base value unavailable; specify a target value explicitly.")
                absolute_default = state.get("absolute")
                if absolute_default is None:
                    absolute_default = base_value if np.isfinite(base_value) else 0.0
                absolute_value = st.number_input(
                    f"Target value ({units})" if units else "Target value",
                    key=f"scenario_absolute_{parameter}",
                    value=float(absolute_default),
                    step=100.0,
                    format="%.2f",
                )
                state["absolute"] = absolute_value
            builder_values[parameter] = state

    st.session_state[SCENARIO_VALUE_STATE_KEY] = builder_values

    scenario_name = st.text_input("Scenario name", key=SCENARIO_NAME_STATE_KEY)
    if st.button("Save Scenario", key="save_scenario_definition"):
        if not scenario_name.strip():
            st.warning("Please provide a scenario name before saving.")
        elif not selected_parameters:
            st.warning("Select at least one parameter to configure.")
        else:
            overrides: Dict[str, float] = {}
            deltas: Dict[str, float | None] = {}
            for parameter in selected_parameters:
                info = catalog_lookup.get(parameter, {})
                base_value = float(info.get("Base Value", np.nan)) if info else np.nan
                state = builder_values.get(parameter, {})
                mode = state.get("mode", "percent")
                if mode == "percent" and np.isfinite(base_value) and not np.isclose(base_value, 0.0):
                    percent = float(state.get("percent", 0.0))
                    target_value = base_value * (1 + percent / 100.0)
                    overrides[parameter] = float(target_value)
                    deltas[parameter] = float(percent)
                else:
                    absolute = state.get("absolute")
                    if absolute is None:
                        absolute = base_value if np.isfinite(base_value) else 0.0
                    absolute = float(absolute)
                    overrides[parameter] = absolute
                    if np.isfinite(base_value) and not np.isclose(base_value, 0.0):
                        change = (absolute / base_value - 1.0) * 100.0
                        deltas[parameter] = float(change)
                    else:
                        deltas[parameter] = None

            new_entry = {
                "name": scenario_name.strip(),
                "overrides": overrides,
                "deltas": deltas,
            }
            existing_index = next(
                (
                    idx
                    for idx, entry in enumerate(scenario_definitions)
                    if entry["name"].casefold() == scenario_name.strip().casefold()
                ),
                None,
            )
            if existing_index is not None:
                scenario_definitions[existing_index] = new_entry
            else:
                scenario_definitions.append(new_entry)
            st.session_state[SCENARIO_DEFINITIONS_KEY] = scenario_definitions
            st.session_state[SCENARIO_CLEAR_NAME_FLAG] = True
            st.success(f"Scenario '{scenario_name}' saved.")
            _trigger_rerun()

    if scenario_definitions:
        st.subheader("Configured Scenarios")
        summary_rows: List[Dict[str, object]] = []
        for entry in scenario_definitions:
            for parameter, target in entry.get("overrides", {}).items():
                delta_value = entry.get("deltas", {}).get(parameter)
                summary_rows.append(
                    {
                        "Scenario": entry["name"],
                        "Parameter": parameter,
                        "Target Value": float(target),
                        "Delta (%)": None if delta_value is None else float(delta_value),
                    }
                )
        summary_df = pd.DataFrame(summary_rows)
        if not summary_df.empty:
            if "Delta (%)" in summary_df.columns:
                summary_df["Delta (%)"] = summary_df["Delta (%)"].round(2)
            st.dataframe(summary_df, use_container_width=True, hide_index=True)
        for idx, entry in enumerate(scenario_definitions):
            if st.button(f"Remove {entry['name']}", key=f"remove_scenario_{idx}"):
                scenario_definitions.pop(idx)
                st.session_state[SCENARIO_DEFINITIONS_KEY] = scenario_definitions
                _trigger_rerun()
                return

    def _scenario_model() -> CassavaBioethanolModel:
        clone = CassavaBioethanolModel(copy.deepcopy(base_page))
        clone.scenario = model.scenario
        return clone

    comparison_model = _scenario_model()
    scenario_configs = [ScenarioConfig(entry["name"], entry["overrides"]) for entry in scenario_definitions]
    scenario_cache: Dict[str, Dict[str, object]] = st.session_state.setdefault(SCENARIO_CACHE_KEY, {})
    cached_entry = scenario_cache.get(model.scenario)
    signature = _scenario_definition_signature(scenario_definitions)
    comparison_button = st.button(
        "Run Scenario Comparison",
        key=f"run_scenario_cmp_{model.scenario.lower()}",
    )

    needs_refresh = (
        comparison_button
        or cached_entry is None
        or cached_entry.get("version") != _current_model_version()
        or cached_entry.get("signature") != signature
    )

    if needs_refresh:
        with st.spinner("Evaluating scenario overrides..."):
            base_result = comparison_model.build()
            base_metrics = base_result.get("metrics", {})
            if scenario_configs:
                comparison_df = scenario_comparison(comparison_model, scenario_configs)
            else:
                comparison_df = pd.DataFrame(columns=["Scenario"])
        scenario_cache[model.scenario] = {
            "version": _current_model_version(),
            "signature": signature,
            "comparison": comparison_df,
            "base_metrics": base_metrics,
        }
        st.session_state[SCENARIO_CACHE_KEY] = scenario_cache
        cached_entry = scenario_cache[model.scenario]

    comparison_df = cached_entry.get("comparison", pd.DataFrame()) if cached_entry else pd.DataFrame()
    base_metrics = cached_entry.get("base_metrics", {}) if cached_entry else {}

    base_row = {"Scenario": "Base Case"}
    base_row.update(base_metrics)
    base_df = pd.DataFrame([base_row])
    comparison_display = comparison_df.copy()
    for metric in ["Project NPV", "Project IRR", "Equity IRR", "Payback Period (years)"]:
        if metric in comparison_display.columns and metric in base_metrics:
            comparison_display[f"{metric} Δ"] = comparison_display[metric] - base_metrics[metric]

    combined_df = pd.concat([base_df, comparison_display], ignore_index=True, sort=False)
    for metric in ["Project NPV", "Project IRR", "Equity IRR", "Payback Period (years)"]:
        delta_col = f"{metric} Δ"
        if delta_col in combined_df.columns:
            combined_df.loc[combined_df["Scenario"] == "Base Case", delta_col] = 0.0

    st.subheader("Scenario Comparison")
    if combined_df.empty:
        st.info("Configure scenarios and click 'Run Scenario Comparison' to evaluate impacts.")
    else:
        preferred_cols = [
            "Scenario",
            "Project NPV",
            "Project NPV Δ",
            "Project IRR",
            "Project IRR Δ",
            "Equity IRR",
            "Equity IRR Δ",
            "Payback Period (years)",
            "Payback Period (years) Δ",
        ]
        display_cols = [col for col in preferred_cols if col in combined_df.columns]
        display_df = combined_df[display_cols] if display_cols else combined_df
        st.dataframe(display_df, use_container_width=True)

    st.subheader("Reverse Stress Test (Correlated)")
    st.caption(
        "Search correlated stress combinations to answer: what breaks DSCR covenant and what breaks NPV > 0."
    )
    if st.button("Run Reverse Stress Test", key=f"reverse_stress_{model.scenario.lower()}"):
        reverse_model = _scenario_model()
        reverse_df = reverse_stress_test(reverse_model, dscr_floor=1.0, npv_floor=0.0)
        if reverse_df.empty:
            st.info("No breach combination found within the configured stress grid.")
        else:
            st.dataframe(reverse_df, use_container_width=True, hide_index=True)

    toolkit = comparison_model.advanced_toolkit()
    analysis_df = pd.DataFrame()
    feature_cols: List[str] = []
    if scenario_definitions and not comparison_df.empty:
        feature_rows: List[Dict[str, object]] = []
        for entry in scenario_definitions:
            row: Dict[str, object] = {"Scenario": entry["name"]}
            for parameter in SCENARIO_PARAMETER_NAMES:
                delta_value = entry.get("deltas", {}).get(parameter)
                row[parameter] = 0.0 if delta_value is None else float(delta_value)
            feature_rows.append(row)
        features_df = pd.DataFrame(feature_rows).set_index("Scenario") if feature_rows else pd.DataFrame()
        metrics_df = comparison_df.set_index("Scenario") if not comparison_df.empty else pd.DataFrame()
        if not features_df.empty and not metrics_df.empty:
            analysis_df = features_df.join(metrics_df, how="inner")
            if "Project NPV" in analysis_df.columns and "Project NPV" in metrics_df.columns:
                base_npv = base_metrics.get("Project NPV")
                if base_npv is not None:
                    analysis_df["Project NPV Change"] = analysis_df["Project NPV"] - base_npv
            if "Equity IRR" in analysis_df.columns and "Equity IRR" in metrics_df.columns:
                base_equity = base_metrics.get("Equity IRR")
                if base_equity is not None:
                    analysis_df["Equity IRR Change"] = analysis_df["Equity IRR"] - base_equity
            feature_cols = [param for param in SCENARIO_PARAMETER_NAMES if param in analysis_df.columns]
            if feature_cols:
                analysis_df[feature_cols] = analysis_df[feature_cols].apply(pd.to_numeric, errors="coerce").fillna(0.0)

    st.subheader("Predictive Scenario Tools")
    regression_tab, tree_tab, time_series_tab, revolver_tab = st.tabs(
        ["Regression", "Decision Tree", "Time Series", "Revolver"]
    )

    with regression_tab:
        if len(analysis_df) >= 2 and feature_cols and "Project NPV Change" in analysis_df.columns:
            regression_input = analysis_df[feature_cols + ["Project NPV Change"]].reset_index(drop=True)
            try:
                regression_result = toolkit.linear_regression(regression_input, "Project NPV Change")
            except Exception as exc:  # pragma: no cover - display feedback only
                st.warning(f"Regression analysis failed: {exc}")
            else:
                metric_cols = st.columns(2)
                metric_cols[0].metric("R²", f"{regression_result.score:.3f}")
                metric_cols[1].metric("Intercept", f"{regression_result.intercept:,.2f}")
                coeff_df = (
                    pd.DataFrame(
                        [
                            {"Parameter": name, "Coefficient": coeff}
                            for name, coeff in regression_result.coefficients.items()
                        ]
                    )
                    .sort_values("Coefficient", key=lambda s: s.abs(), ascending=False)
                )
                st.dataframe(coeff_df, use_container_width=True, hide_index=True)
        else:
            st.info("Add at least two scenarios to run regression analysis.")

    with tree_tab:
        if len(analysis_df) >= 2 and feature_cols and "Project NPV Change" in analysis_df.columns:
            tree_input = analysis_df[feature_cols + ["Project NPV Change"]].reset_index(drop=True)
            try:
                tree_result = toolkit.decision_tree_regression(
                    tree_input,
                    "Project NPV Change",
                    max_depth=min(4, len(feature_cols)),
                )
            except Exception as exc:  # pragma: no cover - display feedback only
                st.warning(f"Decision tree analysis failed: {exc}")
            else:
                metric_cols = st.columns(2)
                score_value = tree_result.score if np.isfinite(tree_result.score) else float("nan")
                metric_cols[0].metric("R²", f"{score_value:.3f}" if np.isfinite(score_value) else "n/a")
                metric_cols[1].metric("Depth", str(tree_result.depth))
                importance_df = (
                    pd.DataFrame(
                        [
                            {"Parameter": name, "Importance": importance}
                            for name, importance in tree_result.feature_importances.items()
                        ]
                    )
                    .sort_values("Importance", ascending=False)
                )
                st.dataframe(importance_df, use_container_width=True, hide_index=True)
        else:
            st.info("Add at least two scenarios to explore decision tree splits.")

    with time_series_tab:
        financials = results.get("financials") if isinstance(results, dict) else None
        annual_cf = getattr(financials, "cashflow_annual", None) if financials is not None else None
        if isinstance(annual_cf, pd.DataFrame) and "Free Cash Flow" in annual_cf.columns:
            series = annual_cf["Free Cash Flow"]
            if isinstance(series, pd.Series) and not series.empty:
                periods = st.slider("Forecast periods", min_value=1, max_value=10, value=5, step=1)
                try:
                    forecast = toolkit.arima_forecast(series, periods=periods)
                except Exception as exc:  # pragma: no cover - display feedback only
                    st.warning(f"Time-series forecast failed: {exc}")
                else:
                    actual_df = series.to_frame(name="Value").reset_index().rename(columns={series.index.name or "index": "Period"})
                    actual_df["Type"] = "Actual"
                    forecast_df = forecast.to_frame(name="Value").reset_index().rename(columns={forecast.index.name or "index": "Period"})
                    forecast_df["Type"] = "Forecast"
                    chart_df = pd.concat([actual_df, forecast_df], ignore_index=True)
                    chart_df["Period"] = chart_df["Period"].astype(str)
                    fig = px.line(chart_df, x="Period", y="Value", color="Type", title="Free Cash Flow Forecast")
                    st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("Free cash flow history is required for time-series forecasting.")
        else:
            st.info("Run the model to generate free cash flow before forecasting.")

    with revolver_tab:
        financials = results.get("financials") if isinstance(results, dict) else None
        monthly_cf = getattr(financials, "cashflow_monthly", None) if financials is not None else None
        if isinstance(monthly_cf, pd.DataFrame) and "Free Cash Flow" in monthly_cf.columns:
            series = monthly_cf["Free Cash Flow"]
            if isinstance(series, pd.Series) and not series.empty:
                window = st.slider("Rolling window (months)", min_value=3, max_value=24, value=12, step=1)
                revolver_df = toolkit.revolver_projection(series, window=window)
                if not revolver_df.empty:
                    viz_df = revolver_df.reset_index().rename(columns={series.index.name or "index": "Period"})
                    viz_df["Period"] = viz_df["Period"].astype(str)
                    melted = viz_df.melt(
                        id_vars=["Period"],
                        value_vars=["Value", "Rolling Mean"],
                        var_name="Measure",
                        value_name="Amount",
                    )
                    fig = px.line(
                        melted,
                        x="Period",
                        y="Amount",
                        color="Measure",
                        title="Rolling Free Cash Flow (Revolver)",
                    )
                    st.plotly_chart(fig, use_container_width=True)
                    st.caption("Rolling standard deviation provides a volatility view of the revolver balance.")
                    std_chart = viz_df.set_index("Period")["Rolling Std"].rename("Rolling Std Dev")
                    st.line_chart(std_chart)
            else:
                st.info("Monthly free cash flow is required to analyse the revolver trajectory.")
        else:
            st.info("Run the model to generate monthly free cash flow before analysing the revolver.")

    st.subheader("Goal Seek Results")
    goal_seek_parameter = "Corporate tax rate"
    goal_seek_metric = "Project NPV"
    goal_message: str | None = None
    empty_goal_df = pd.DataFrame(
        columns=[
            "Parameter",
            "Target Metric",
            "Target Value",
            "Target Name",
            "Achieved Value",
            "Tolerance",
            "Iterations",
        ]
    )

    try:
        target_value = float(results["metrics"].get(goal_seek_metric, 0.0))
        goal_model = _scenario_model()
        goal_result = goal_seek_to_target(goal_model, goal_seek_parameter, goal_seek_metric, target_value)
        goal_df = pd.DataFrame(
            [
                {
                    "Parameter": goal_seek_parameter,
                    "Target Metric": goal_seek_metric,
                    "Target Value": target_value,
                    "Target Name": goal_result.target_name,
                    "Achieved Value": goal_result.achieved_value,
                    "Tolerance": goal_result.tolerance,
                    "Iterations": goal_result.iterations,
                }
            ]
        )
    except KeyError:
        goal_df = empty_goal_df
        goal_message = "The selected goal seek parameter is not available in the current global inputs."
    except ValueError as exc:
        goal_df = empty_goal_df
        goal_message = str(exc)

    if goal_message:
        st.info(goal_message)

    st.dataframe(goal_df, use_container_width=True)



RAG_PROVIDER_OPTIONS = ["OpenAI", "Copilot", "Vertex", "Anthropic", "DeepSeek", "Custom"]


def _rag_state() -> Dict[str, object]:
    return st.session_state.setdefault(
        "rag_assistant",
        {
            "indexed_docs": [],
            "chunks": [],
            "config": {},
            "insights": "",
            "business_plan": "",
            "forecast_table": pd.DataFrame(),
        },
    )


def _chunk_text(text: str, chunk_size: int = 1400) -> List[str]:
    cleaned = re.sub(r"\s+", " ", text or "").strip()
    if not cleaned:
        return []
    return [cleaned[i : i + chunk_size] for i in range(0, len(cleaned), chunk_size)]


def _build_model_knowledge_pack(results: Dict[str, object]) -> str:
    metrics = results.get("metrics", {}) if isinstance(results, dict) else {}
    financials = results.get("financials") if isinstance(results, dict) else None
    sections: List[str] = ["# Financial Model Knowledge Pack"]

    if metrics:
        metric_rows = [f"- {k}: {v}" for k, v in list(metrics.items())[:25]]
        sections.append("## Core Metrics\n" + "\n".join(metric_rows))

    def _df_section(title: str, df: pd.DataFrame | None, rows: int = 24) -> None:
        if isinstance(df, pd.DataFrame) and not df.empty:
            sections.append(f"## {title}\n" + df.head(rows).to_markdown())

    if financials is not None:
        _df_section("Income Statement (Monthly excerpt)", getattr(financials, "income_monthly", None))
        _df_section("Cash Flow Statement (Monthly excerpt)", getattr(financials, "cashflow_monthly", None))
        _df_section("Balance Sheet (Monthly excerpt)", getattr(financials, "balance_monthly", None))

    for key in ("production", "revenue", "break_even", "payback"):
        payload = results.get(key) if isinstance(results, dict) else None
        if isinstance(payload, pd.DataFrame):
            _df_section(key.replace("_", " ").title(), payload)
        elif hasattr(payload, "monthly") and isinstance(payload.monthly, pd.DataFrame):
            _df_section(f"{key.title()} Monthly", payload.monthly)
        elif hasattr(payload, "annual") and isinstance(payload.annual, pd.DataFrame):
            _df_section(f"{key.title()} Annual", payload.annual)

    return "\n\n".join(sections)


def _simple_retrieve(chunks: List[str], question: str, top_k: int = 6) -> List[str]:
    if not chunks:
        return []
    terms = [t for t in re.findall(r"[a-zA-Z0-9]+", (question or "").lower()) if len(t) > 2]
    if not terms:
        return chunks[:top_k]
    scored = []
    for chunk in chunks:
        lower = chunk.lower()
        score = sum(lower.count(term) for term in terms)
        scored.append((score, chunk))
    ranked = [c for s, c in sorted(scored, key=lambda x: x[0], reverse=True) if s > 0]
    return (ranked or chunks)[:top_k]


def _build_forecast(results: Dict[str, object], years: int) -> pd.DataFrame:
    financials = results.get("financials") if isinstance(results, dict) else None
    if financials is None:
        return pd.DataFrame()
    annual = getattr(financials, "income_annual", pd.DataFrame())
    if annual.empty or "Revenue" not in annual.columns:
        return pd.DataFrame()
    annual = annual.copy()
    if isinstance(annual.index, pd.DatetimeIndex):
        idx = annual.index
    elif isinstance(annual.index, pd.PeriodIndex):
        idx = annual.index.to_timestamp()
    else:
        raw_index = pd.Index(annual.index)
        numeric_years = pd.to_numeric(raw_index, errors="coerce")
        if numeric_years.notna().all():
            idx = pd.to_datetime(numeric_years.astype(int).astype(str), format="%Y", errors="coerce")
        else:
            idx = pd.to_datetime(raw_index, errors="coerce")
    annual.index = idx
    annual = annual[~annual.index.isna()]
    if annual.empty:
        return pd.DataFrame()

    rev = pd.to_numeric(annual["Revenue"], errors="coerce").dropna()
    ebitda = pd.to_numeric(annual.get("EBITDA"), errors="coerce").dropna() if "EBITDA" in annual.columns else pd.Series(dtype=float)
    if rev.empty:
        return pd.DataFrame()

    if len(rev) > 1 and rev.iloc[0] > 0:
        cagr = (rev.iloc[-1] / rev.iloc[0]) ** (1 / (len(rev) - 1)) - 1
    else:
        cagr = 0.03

    last_year = int(rev.index[-1].year)
    last_rev = float(rev.iloc[-1])
    last_ebitda = float(ebitda.iloc[-1]) if not ebitda.empty else 0.0
    margin = (last_ebitda / last_rev) if last_rev else 0.0

    rows = []
    cur_rev = last_rev
    for i in range(1, max(1, years) + 1):
        cur_rev *= 1 + cagr
        rows.append({
            "Year": last_year + i,
            "Forecast Revenue": cur_rev,
            "Forecast EBITDA": cur_rev * margin,
            "Assumed Growth": cagr,
        })
    return pd.DataFrame(rows)




def _round_nearest_100(df: pd.DataFrame | None) -> pd.DataFrame:
    """Return a copy with numeric financial figures rounded to nearest 100.

    Calendar/index columns such as ``Year`` are intentionally excluded.
    """

    if not isinstance(df, pd.DataFrame):
        return pd.DataFrame()

    rounded = df.copy()
    numeric_cols = list(rounded.select_dtypes(include=[np.number]).columns)
    if not numeric_cols:
        return rounded

    protected = {
        c
        for c in numeric_cols
        if str(c).strip().lower() in {"year", "month", "date", "period", "start year", "end year"}
    }

    for col in numeric_cols:
        if col in protected:
            continue
        rounded[col] = (pd.to_numeric(rounded[col], errors="coerce") / 100.0).round() * 100.0

    return rounded


def _to_markdown_table(df: pd.DataFrame | None, rows: int = 20) -> str:
    if not isinstance(df, pd.DataFrame) or df.empty:
        return "_No data available._"
    view = _round_nearest_100(df).head(rows).copy()
    if isinstance(view.index, pd.DatetimeIndex):
        view.index = view.index.to_period("Y").astype(str)
    try:
        return view.to_markdown()
    except ImportError:
        # `to_markdown` requires optional `tabulate`; provide a robust fallback.
        return "```\n" + view.to_string() + "\n```"


def _metric_commentary(metrics: Dict[str, object]) -> str:
    irr = _annualise(metrics.get("Project IRR"))
    eq_irr = _annualise(metrics.get("Equity IRR"))
    payback = metrics.get("Payback Period (years)")
    dscr = metrics.get("DSCR (min)")
    llcr = metrics.get("LLCR")
    plcr = metrics.get("PLCR")
    return (
        "- **Valuation**: Project NPV reflects value creation potential under current assumptions.\n"
        f"- **Returns**: Project IRR is approximately {_format_rate(irr)} and Equity IRR is {_format_rate(eq_irr)}.\n"
        f"- **Liquidity recovery**: Payback is {payback if payback is not None else 'n/a'} years.\n"
        f"- **Debt service strength**: Minimum DSCR is {_format_rate(dscr)} with LLCR {_format_rate(llcr)} and PLCR {_format_rate(plcr)}."
    )


def _compose_business_plan(
    results: Dict[str, object],
    insights: str,
    years: int,
    forecast_df: pd.DataFrame,
    frames: Dict[str, pd.DataFrame] | None = None,
) -> str:
    """Build a professional investor-style business plan text package."""

    financials = results.get("financials")
    metrics = results.get("metrics", {}) if isinstance(results, dict) else {}

    income_annual = _round_nearest_100(getattr(financials, "income_annual", pd.DataFrame())) if financials is not None else pd.DataFrame()
    cashflow_annual = _round_nearest_100(getattr(financials, "cashflow_annual", pd.DataFrame())) if financials is not None else pd.DataFrame()
    balance_annual = _round_nearest_100(getattr(financials, "balance_annual", pd.DataFrame())) if financials is not None else pd.DataFrame()

    production = results.get("production")
    production_annual = _round_nearest_100(getattr(production, "annual", pd.DataFrame())) if production is not None else pd.DataFrame()
    revenue = results.get("revenue")
    revenue_annual = _round_nearest_100(getattr(revenue, "annual", pd.DataFrame())) if revenue is not None else pd.DataFrame()

    metric_table = _round_nearest_100(pd.DataFrame([metrics]).T.reset_index())
    metric_table.columns = ["Key Metric", "Value"]

    income_plot_cols = [c for c in ["Revenue", "EBITDA", "Net Income"] if c in income_annual.columns]
    cash_plot_cols = [c for c in ["Operating Cash Flow", "Investing Cash Flow", "Financing Cash Flow", "Free Cash Flow"] if c in cashflow_annual.columns]
    balance_plot_cols = [c for c in ["Total Assets", "Total Liabilities & Equity", "Debt", "Equity"] if c in balance_annual.columns]

    income_plot_df = income_annual[income_plot_cols].copy() if income_plot_cols else pd.DataFrame()
    cash_plot_df = cashflow_annual[cash_plot_cols].copy() if cash_plot_cols else pd.DataFrame()
    balance_plot_df = balance_annual[balance_plot_cols].copy() if balance_plot_cols else pd.DataFrame()

    scenario = metrics.get("Scenario", "FARM_ONLY")
    dscr_min = _format_rate(metrics.get("DSCR (min)"))
    llcr = _format_rate(metrics.get("LLCR"))
    plcr = _format_rate(metrics.get("PLCR"))

    sensitivity_df = (frames or {}).get("Sensitivity Analyses", pd.DataFrame())
    scenario_df = (frames or {}).get("Scenario / IFs Analysis", pd.DataFrame())
    monte_carlo_df = (frames or {}).get("Monte Carlo Simulation", pd.DataFrame())

    return f"""# Cassava Bioethanol Comprehensive Business Plan

## 1. Executive Summary
This plan is generated from the integrated cassava-ethanol financial model and RAG evidence base for scenario **{scenario}**.
The current investment thesis is supported by diversified revenue (fuel ethanol + animal feed), integrated debt-service diagnostics, and projection-consistent forecasting.

## 2. Investment Highlights and Key Metrics
### Professional Interpretation
{_metric_commentary(metrics)}

### Key Metrics Dashboard
Interpretation: this dashboard is the decision core for equity return quality, debt-service resilience, and valuation headroom.
{_to_markdown_table(metric_table, rows=50)}

## 3. Market, Commercial Strategy, and Operations
### 3.1 Commercial Positioning
The model embeds offtake floor/ceiling assumptions, take-or-pay coverage, and contracted feedstock mechanisms to reflect realistic contract economics.

### 3.2 Operational Configuration
Production planning, capex, staffing, opex, and working-capital assumptions are integrated into three-statement outputs and debt metrics.

## 4. Annual Financial Statements (Reproduced)
### 4.1 Annual Income Statement
Interpretation: this section explains top-line growth, operating profitability, and net earnings conversion quality.
{_to_markdown_table(income_annual, rows=25)}

### 4.2 Annual Cash Flow Statement
Interpretation: this section tracks cash generation, reinvestment intensity, and financing dependence over time.
{_to_markdown_table(cashflow_annual, rows=25)}

### 4.3 Annual Balance Sheet
Interpretation: this section highlights capital structure strength, leverage trajectory, and net asset accumulation.
{_to_markdown_table(balance_annual, rows=25)}

## 5. Schedules and Forecasts
### 5.1 Production Annual Schedule
{_to_markdown_table(production_annual, rows=25)}

### 5.2 Revenue Annual Schedule
{_to_markdown_table(revenue_annual, rows=25)}

### 5.3 Forecast (Projection-Horizon Aligned)
Forecast horizon: **{years} year(s)**.
{_to_markdown_table(forecast_df, rows=years + 2)}

## 6. Graphs and Plot Data (Reproduced for Download Pack)
The following data tables are the source series for the charts included in the download package.

### 6.1 Income Statement Plot Data
{_to_markdown_table(income_plot_df, rows=25)}

### 6.2 Cash Flow Plot Data
{_to_markdown_table(cash_plot_df, rows=25)}

### 6.3 Balance Sheet Plot Data
{_to_markdown_table(balance_plot_df, rows=25)}

### 6.4 Forecast Plot Data
{_to_markdown_table(forecast_df, rows=years + 2)}

## 7. Lender/Covenant Narrative
Minimum DSCR is **{dscr_min}** with LLCR **{llcr}** and PLCR **{plcr}**. These indicators frame debt-service resilience and refinancing feasibility under the modeled assumptions.

## 8. Financial Performance
Interpretation: financial performance combines revenue scale-up, EBITDA quality, and bottom-line durability.
{_to_markdown_table(income_annual, rows=25)}

## 9. Financial Position
Interpretation: financial position demonstrates solvency, leverage profile, and balance-sheet risk absorption capacity.
{_to_markdown_table(balance_annual, rows=25)}

## 10. Cash Flow Statement
Interpretation: cash flow analysis validates debt serviceability and reinvestment capacity under operating assumptions.
{_to_markdown_table(cashflow_annual, rows=25)}

## 11. Sensitivity Analyses
Interpretation: sensitivity analyses quantify first-order valuation and return responses to key assumption shocks.
{_to_markdown_table(sensitivity_df, rows=40)}

## 12. Scenario / IFs Analysis
Interpretation: scenario analysis compares strategic configurations and downside/upside outcomes across selected cases.
{_to_markdown_table(scenario_df, rows=40)}

## 13. Monte Carlo Simulation
Interpretation: Monte Carlo outputs characterize distributional risk and confidence intervals for key investment outcomes.
{_to_markdown_table(monte_carlo_df, rows=40)}

## 14. Risk and Mitigation
Risk register impacts, commercial safeguards, and scenario analytics are embedded to produce risk-adjusted performance views and improve investor decision confidence.

## 15. AI/RAG Supporting Insights
{insights or '_No additional RAG insights available. Run Document Indexing and Run the AI._'}

## 16. Conclusion and Funding Case
The model output indicates a financeable platform when contract quality, feedstock reliability, and covenant headroom are maintained. Recommended next step is lender term-sheet calibration against DSCR/LLCR/PLCR constraints and downside scenarios.
"""



def _rag_export_frames(model: CassavaBioethanolModel, results: Dict[str, object], forecast_df: pd.DataFrame) -> Dict[str, pd.DataFrame]:
    financials = results.get("financials") if isinstance(results, dict) else None
    metrics = results.get("metrics", {}) if isinstance(results, dict) else {}

    income_annual = _round_nearest_100(getattr(financials, "income_annual", pd.DataFrame())) if financials is not None else pd.DataFrame()
    cashflow_annual = _round_nearest_100(getattr(financials, "cashflow_annual", pd.DataFrame())) if financials is not None else pd.DataFrame()
    balance_annual = _round_nearest_100(getattr(financials, "balance_annual", pd.DataFrame())) if financials is not None else pd.DataFrame()
    forecast = _round_nearest_100(forecast_df)

    metrics_df = pd.DataFrame(list(metrics.items()), columns=["Key Metric", "Value"])
    if not metrics_df.empty:
        metrics_df["Value"] = pd.to_numeric(metrics_df["Value"], errors="ignore")
        metrics_df = _round_nearest_100(metrics_df)

    sensitivity_df = pd.DataFrame()
    try:
        base_page = copy.deepcopy(results.get("input_page_snapshot", model.input_page))
        sensitivity_model = CassavaBioethanolModel(copy.deepcopy(base_page))
        sensitivity_model.scenario = model.scenario
        if DEFAULT_SENSITIVITY_SCENARIOS:
            sensitivity_df = _round_nearest_100(run_sensitivity(sensitivity_model, DEFAULT_SENSITIVITY_SCENARIOS))
    except Exception:
        sensitivity_df = pd.DataFrame()

    scenario_rows: List[Dict[str, object]] = []
    try:
        snapshot = copy.deepcopy(results.get("input_page_snapshot", model.input_page))
        for scenario_name in CassavaBioethanolModel.SCENARIOS:
            scenario_model = CassavaBioethanolModel(copy.deepcopy(snapshot))
            scenario_model.scenario = scenario_name
            scenario_result = scenario_model.build(scenario_name)
            scenario_metrics = scenario_result.get("metrics", {}) if isinstance(scenario_result, dict) else {}
            scenario_rows.append(
                {
                    "Scenario": scenario_name,
                    "Project NPV": scenario_metrics.get("Project NPV"),
                    "Project IRR": scenario_metrics.get("Project IRR"),
                    "Equity IRR": scenario_metrics.get("Equity IRR"),
                    "Payback Period (years)": scenario_metrics.get("Payback Period (years)"),
                    "DSCR (min)": scenario_metrics.get("DSCR (min)"),
                }
            )
    except Exception:
        scenario_rows = []
    scenario_df = _round_nearest_100(pd.DataFrame(scenario_rows))

    monte_carlo_summary = pd.DataFrame()
    try:
        cache: Dict[str, Dict[str, object]] = st.session_state.get(MC_CACHE_KEY, {})
        cached_entry = cache.get(model.scenario, {})
        mc_results = cached_entry.get("results")
        if isinstance(mc_results, pd.DataFrame) and not mc_results.empty:
            monte_carlo_summary = mc_results.describe(percentiles=[0.1, 0.25, 0.5, 0.75, 0.9]).T.reset_index()
            monte_carlo_summary = monte_carlo_summary.rename(columns={"index": "Metric"})
    except Exception:
        monte_carlo_summary = pd.DataFrame()
    monte_carlo_summary = _round_nearest_100(monte_carlo_summary)

    return {
        "Key Metrics": metrics_df,
        "Key Metrics Dashboard": metrics_df,
        "Financial Performance": income_annual,
        "Financial Position": balance_annual,
        "Cash Flow Statement": cashflow_annual,
        "Income Annual": income_annual,
        "Cash Flow Annual": cashflow_annual,
        "Balance Annual": balance_annual,
        "Forecast": forecast,
        "Sensitivity Analyses": sensitivity_df,
        "Scenario / IFs Analysis": scenario_df,
        "Monte Carlo Simulation": monte_carlo_summary,
    }


def _prepare_export_table(df: pd.DataFrame, max_rows: int = 40) -> pd.DataFrame:
    """Normalize dataframe for export readability (clear rows/columns)."""
    if not isinstance(df, pd.DataFrame) or df.empty:
        return pd.DataFrame()
    view = df.copy()
    if not isinstance(view.index, pd.RangeIndex) or view.index.name:
        view = view.reset_index()
    view.columns = [str(c) for c in view.columns]
    return view.head(max_rows)


def _fit_column_widths_excel(ws, df: pd.DataFrame, max_width: int = 28) -> None:
    if not isinstance(df, pd.DataFrame) or df.empty:
        return
    for idx, col in enumerate(df.columns):
        sample = [str(col)] + [str(v) for v in df[col].head(30).tolist()]
        width = min(max(len(x) for x in sample) + 2, max_width)
        ws.set_column(idx, idx, width)


def _generate_word_business_plan_bytes(plan_text: str, frames: Dict[str, pd.DataFrame]) -> bytes:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Inches, Pt
    except Exception:
        return plan_text.encode("utf-8")

    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    title = doc.add_heading("Cassava Bioethanol Business Plan", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for line in (plan_text or "").splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("# "):
            continue
        elif stripped:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(4)
        else:
            doc.add_paragraph("")

    section_notes = {
        "Key Metrics": "Interpretation: this dashboard summarizes project returns, valuation, and covenant readiness.",
        "Key Metrics Dashboard": "Interpretation: this dashboard summarizes project returns, valuation, and covenant readiness.",
        "Financial Performance": "Interpretation: this table highlights revenue, profitability, and margin dynamics.",
        "Financial Position": "Interpretation: this table shows solvency, leverage, and capital structure progression.",
        "Cash Flow Statement": "Interpretation: this table captures operating, investing, and financing cash dynamics.",
        "Sensitivity Analyses": "Interpretation: this section quantifies how outputs move when key assumptions are stressed.",
        "Scenario / IFs Analysis": "Interpretation: this section compares strategic alternatives against the base case.",
        "Monte Carlo Simulation": "Interpretation: this section summarizes distribution-based risk and confidence bands.",
    }

    def _add_table(title: str, df: pd.DataFrame) -> None:
        view = _prepare_export_table(df, max_rows=50)
        if view.empty:
            return
        doc.add_heading(title, level=2)
        if title in section_notes:
            doc.add_paragraph(section_notes[title])
        table = doc.add_table(rows=1, cols=len(view.columns))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, col in enumerate(view.columns):
            run = hdr_cells[i].paragraphs[0].add_run(str(col))
            run.bold = True

        for _, row in view.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row.tolist()):
                cells[i].text = str(val)

        doc.add_paragraph("")

    for title, frame in frames.items():
        _add_table(title, frame)

    def _add_chart(title: str, df: pd.DataFrame) -> None:
        view = _prepare_export_table(df, max_rows=40)
        if view.empty or view.shape[1] < 2:
            return
        fig, ax = plt.subplots(figsize=(10.5, 3.8))
        x = view.iloc[:, 0].astype(str)
        for col in view.columns[1: min(6, len(view.columns))]:
            ax.plot(x, pd.to_numeric(view[col], errors="coerce"), label=str(col), linewidth=2)
        ax.set_title(title)
        ax.tick_params(axis="x", rotation=45)
        ax.grid(True, alpha=0.2)
        ax.legend(loc="best", fontsize=8)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=170)
        plt.close(fig)
        buf.seek(0)
        doc.add_picture(buf, width=Inches(9.8))
        doc.add_paragraph("")

    for title, frame in frames.items():
        _add_chart(f"{title} Trend", frame)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _generate_pdf_business_plan_bytes(plan_text: str, frames: Dict[str, pd.DataFrame]) -> bytes:
    out = io.BytesIO()
    section_notes = {
        "Key Metrics": "Interpretation: dashboard for returns, value, and covenant quality.",
        "Key Metrics Dashboard": "Interpretation: dashboard for returns, value, and covenant quality.",
        "Financial Performance": "Interpretation: trend in revenue, EBITDA and net income quality.",
        "Financial Position": "Interpretation: leverage and balance-sheet resilience over time.",
        "Cash Flow Statement": "Interpretation: cash generation and debt serviceability profile.",
        "Sensitivity Analyses": "Interpretation: impact of assumption shifts on key outputs.",
        "Scenario / IFs Analysis": "Interpretation: strategic-case comparison versus base case.",
        "Monte Carlo Simulation": "Interpretation: probabilistic distribution of outcomes.",
    }
    with PdfPages(out) as pdf:
        fig, ax = plt.subplots(figsize=(11.69, 8.27))
        ax.axis("off")
        wrapped = textwrap.fill((plan_text or "").replace("\n", " "), width=130)
        ax.text(0.01, 0.99, wrapped[:14000], va="top", ha="left", fontsize=8)
        pdf.savefig(fig, bbox_inches="tight")
        plt.close(fig)

        for title, df in frames.items():
            view = _prepare_export_table(df, max_rows=30)
            if view.empty:
                continue

            fig_t, ax_t = plt.subplots(figsize=(11.69, 8.27))
            ax_t.axis("off")
            ax_t.set_title(title, fontsize=14, pad=14)
            note = section_notes.get(title)
            if note:
                ax_t.text(0.01, 0.95, note, transform=ax_t.transAxes, fontsize=9, ha="left", va="top")
            tbl = ax_t.table(
                cellText=view.astype(str).values,
                colLabels=[str(c) for c in view.columns],
                loc="center",
                cellLoc="center",
            )
            tbl.auto_set_font_size(False)
            tbl.set_fontsize(7.5)
            tbl.scale(1.2, 1.25)
            pdf.savefig(fig_t, bbox_inches="tight")
            plt.close(fig_t)

            if view.shape[1] >= 2:
                fig_c, ax_c = plt.subplots(figsize=(11.69, 4.6))
                x = view.iloc[:, 0].astype(str)
                for col in view.columns[1: min(6, len(view.columns))]:
                    ax_c.plot(x, pd.to_numeric(view[col], errors="coerce"), label=str(col), linewidth=2)
                ax_c.set_title(f"{title} Trend")
                ax_c.tick_params(axis="x", rotation=45)
                ax_c.grid(True, alpha=0.25)
                ax_c.legend(loc="best", fontsize=8)
                plt.tight_layout()
                pdf.savefig(fig_c, bbox_inches="tight")
                plt.close(fig_c)
    out.seek(0)
    return out.getvalue()


def _render_rag_assistant_page(model: CassavaBioethanolModel, results: Dict[str, object]) -> None:
    st.subheader("RAG Assistant")
    st.caption("Upload reference materials, index model outputs, and generate a comprehensive business-plan draft with forecasts.")

    rag = _rag_state()
    left, right = st.columns([1.05, 1.25])

    with left:
        st.markdown("### 1) Upload reference document")
        uploads = st.file_uploader(
            "Upload business plan references",
            type=["txt", "md", "pdf", "docx", "csv", "xlsx"],
            accept_multiple_files=True,
            key="rag_uploads",
        )

        st.markdown("### 2) AI & Machine Learning Configuration")
        provider = st.selectbox("AI options", RAG_PROVIDER_OPTIONS, key="rag_provider")
        custom_provider = st.text_input("Custom AI provider", key="rag_provider_custom") if provider == "Custom" else ""
        model_name = st.text_input("AI model space", value="gpt-4.1", key="rag_model")
        snapshot = results.get("input_page_snapshot") if isinstance(results, dict) else None
        if isinstance(snapshot, InputLandingPage):
            projection_years = int(snapshot.projection.end_year) - int(snapshot.projection.start_year) + 1
        else:
            projection_years = int(model.input_page.projection.end_year) - int(model.input_page.projection.start_year) + 1
        projection_years = max(1, projection_years)
        forecast_years = st.number_input(
            "Forecast years",
            min_value=projection_years,
            max_value=projection_years,
            value=projection_years,
            step=1,
            key="rag_forecast_years",
            help="Forecast years are locked to match the Projection Horizon.",
        )
        api_key = st.text_input("API Key", type="password", key="rag_api_key")
        advanced_features = st.multiselect(
            "Generative Features",
            [
                "Auto executive summary",
                "Risk heatmap narrative",
                "Covenant commentary",
                "Capex/opex optimisation ideas",
                "Scenario narrative builder",
                "Investor Q&A prep",
            ],
            default=["Auto executive summary", "Investor Q&A prep"],
            key="rag_features",
        )

        if st.button("Save configuration", key="rag_save_config"):
            rag["config"] = {
                "provider": custom_provider.strip() if provider == "Custom" else provider,
                "model": model_name.strip(),
                "forecast_years": int(forecast_years),
                "api_key_present": bool(api_key),
                "features": advanced_features,
                "saved_at": datetime.utcnow().isoformat(timespec="seconds"),
            }
            st.success("Configuration saved.")

    with right:
        st.markdown("### 3) AI Insights")
        action_cols = st.columns(3)
        with action_cols[0]:
            run_ai = st.button("Run the AI", key="rag_run_ai")
        with action_cols[1]:
            index_btn = st.button("Document Indexing", key="rag_index")
        with action_cols[2]:
            clear_btn = st.button("Clear document indexed", key="rag_clear")

        if clear_btn:
            rag["indexed_docs"] = []
            rag["chunks"] = []
            rag["insights"] = ""
            rag["business_plan"] = ""
            rag["forecast_table"] = pd.DataFrame()
            st.success("Indexed corpus cleared.")

        if index_btn:
            chunks: List[str] = []
            indexed_docs: List[str] = []
            for f in uploads or []:
                content = ""
                try:
                    if f.type in {"text/plain", "text/markdown", "text/csv"} or f.name.lower().endswith((".txt", ".md", ".csv")):
                        content = f.getvalue().decode("utf-8", errors="ignore")
                    else:
                        content = f"Document file: {f.name} (binary file indexed by name only in this local mode)."
                except Exception:
                    content = f"Document file: {f.name} (unable to parse text in local mode)."
                doc_chunks = _chunk_text(content)
                chunks.extend(doc_chunks if doc_chunks else [content])
                indexed_docs.append(f.name)

            chunks.extend(_chunk_text(_build_model_knowledge_pack(results)))
            rag["chunks"] = chunks
            rag["indexed_docs"] = indexed_docs
            st.success(f"Indexed {len(indexed_docs)} uploaded documents and synced model outputs ({len(chunks)} chunks).")

        if rag.get("indexed_docs"):
            st.write("Indexed documents:", ", ".join(rag["indexed_docs"]))

        question = st.text_area("Ask a question", key="rag_question", height=90)
        if run_ai:
            retrieved = _simple_retrieve(rag.get("chunks", []), question or "business plan summary")
            rag["insights"] = "\n\n".join(retrieved) if retrieved else "No indexed content yet. Run Document Indexing first."
            st.success("AI insights generated.")

        if rag.get("insights"):
            st.markdown("#### Retrieved AI insights")
            st.write(rag["insights"][:8000])

    st.markdown("### 4) Prepare Business Plan")
    if st.button("Prepare Business Plan", type="primary", key="rag_prepare_plan"):
        snapshot = results.get("input_page_snapshot") if isinstance(results, dict) else None
        if isinstance(snapshot, InputLandingPage):
            years = int(snapshot.projection.end_year) - int(snapshot.projection.start_year) + 1
        else:
            years = int(model.input_page.projection.end_year) - int(model.input_page.projection.start_year) + 1
        years = max(1, years)
        forecast_df = _round_nearest_100(_build_forecast(results, years))
        rag["forecast_table"] = forecast_df
        narrative = rag.get("insights") or ""
        export_frames = _rag_export_frames(model, results, forecast_df)
        rag["export_frames"] = export_frames
        rag["business_plan"] = _compose_business_plan(results, narrative, years, forecast_df, export_frames)
        st.success("Business plan draft prepared with annual financial tables, professional write-ups, and chart coverage.")

    forecast_df = rag.get("forecast_table", pd.DataFrame())

    st.info("Prepare Business Plan outputs (forecast tables, annual financials, graphs, and plot datasets) are provided in the downloadable Word, PDF, and Excel files only.")

    st.markdown("### 6) Business Plan Downloads")
    plan_text = rag.get("business_plan", "")
    if plan_text:
        frames = rag.get("export_frames") if isinstance(rag.get("export_frames"), dict) else _rag_export_frames(model, results, forecast_df)
        word_bytes = _generate_word_business_plan_bytes(plan_text, frames)
        pdf_bytes = _generate_pdf_business_plan_bytes(plan_text, frames)

        word_name = "Business_Plan.docx" if word_bytes[:2] == b"PK" else "Business_Plan.txt"
        word_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if word_name.endswith(".docx") else "text/plain"
        st.download_button("Download Business Plan (Word)", data=word_bytes, file_name=word_name, mime=word_mime)
        excel_buf = io.BytesIO()
        with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
            if isinstance(forecast_df, pd.DataFrame) and not forecast_df.empty:
                forecast_df.to_excel(writer, sheet_name="Forecast", index=False)
            metrics_df = pd.DataFrame(list(results.get("metrics", {}).items()), columns=["Key Metric", "Value"])
            metrics_df["Value"] = pd.to_numeric(metrics_df["Value"], errors="ignore")
            _round_nearest_100(metrics_df).to_excel(writer, sheet_name="Metrics", index=False)
            fin = results.get("financials") if isinstance(results, dict) else None
            if fin is not None:
                income_annual = getattr(fin, "income_annual", pd.DataFrame())
                cashflow_annual = getattr(fin, "cashflow_annual", pd.DataFrame())
                balance_annual = getattr(fin, "balance_annual", pd.DataFrame())

                if isinstance(income_annual, pd.DataFrame) and not income_annual.empty:
                    income_view = _round_nearest_100(income_annual).reset_index()
                    income_view.to_excel(writer, sheet_name="Income_Annual", index=False)
                else:
                    income_view = pd.DataFrame()

                if isinstance(cashflow_annual, pd.DataFrame) and not cashflow_annual.empty:
                    cash_view = _round_nearest_100(cashflow_annual).reset_index()
                    cash_view.to_excel(writer, sheet_name="Cashflow_Annual", index=False)
                else:
                    cash_view = pd.DataFrame()

                if isinstance(balance_annual, pd.DataFrame) and not balance_annual.empty:
                    balance_view = _round_nearest_100(balance_annual).reset_index()
                    balance_view.to_excel(writer, sheet_name="Balance_Annual", index=False)
                else:
                    balance_view = pd.DataFrame()

                # Plot data sheets to keep the same graph series across Word/PDF/Excel artifacts.
                income_cols = [c for c in ["Revenue", "EBITDA", "Net Income"] if c in getattr(income_annual, 'columns', [])]
                cash_cols = [c for c in ["Operating Cash Flow", "Investing Cash Flow", "Financing Cash Flow", "Free Cash Flow"] if c in getattr(cashflow_annual, 'columns', [])]
                balance_cols = [c for c in ["Total Assets", "Total Liabilities & Equity", "Debt", "Equity"] if c in getattr(balance_annual, 'columns', [])]

                income_plot = _round_nearest_100(income_annual[income_cols]).reset_index() if income_cols else pd.DataFrame()
                cash_plot = _round_nearest_100(cashflow_annual[cash_cols]).reset_index() if cash_cols else pd.DataFrame()
                balance_plot = _round_nearest_100(balance_annual[balance_cols]).reset_index() if balance_cols else pd.DataFrame()

                if not income_plot.empty:
                    income_plot.to_excel(writer, sheet_name="Plot_Income", index=False)
                if not cash_plot.empty:
                    cash_plot.to_excel(writer, sheet_name="Plot_Cashflow", index=False)
                if not balance_plot.empty:
                    balance_plot.to_excel(writer, sheet_name="Plot_Balance", index=False)
                if isinstance(forecast_df, pd.DataFrame) and not forecast_df.empty:
                    _round_nearest_100(forecast_df).to_excel(writer, sheet_name="Plot_Forecast", index=False)

                workbook = writer.book

                def _add_line_chart(sheet_name: str, title: str, max_series: int = 4) -> None:
                    if sheet_name not in writer.sheets:
                        return
                    ws = writer.sheets[sheet_name]
                    # infer dimensions from worksheet write range using dataframe shape fallback
                    df_map = {
                        "Plot_Income": income_plot,
                        "Plot_Cashflow": cash_plot,
                        "Plot_Balance": balance_plot,
                        "Plot_Forecast": forecast_df if isinstance(forecast_df, pd.DataFrame) else pd.DataFrame(),
                    }
                    df = df_map.get(sheet_name, pd.DataFrame())
                    if df.empty or df.shape[1] < 2:
                        return
                    chart = workbook.add_chart({"type": "line"})
                    rows = len(df)
                    cols = min(df.shape[1] - 1, max_series)
                    for idx in range(1, cols + 1):
                        chart.add_series({
                            "name":       [sheet_name, 0, idx],
                            "categories": [sheet_name, 1, 0, rows, 0],
                            "values":     [sheet_name, 1, idx, rows, idx],
                        })
                    chart.set_title({"name": title})
                    chart.set_legend({"position": "bottom"})
                    ws.insert_chart('H2', chart)

                def _format_sheet(sheet_name: str, df: pd.DataFrame | None = None) -> None:
                    ws = writer.sheets.get(sheet_name)
                    if ws is None:
                        return
                    ws.set_landscape()
                    ws.fit_to_pages(1, 0)
                    ws.set_zoom(95)
                    ws.freeze_panes(1, 1)
                    ws.autofilter(0, 0, 0, max((len(df.columns) - 1), 0) if isinstance(df, pd.DataFrame) else 0)
                    if isinstance(df, pd.DataFrame) and not df.empty:
                        _fit_column_widths_excel(ws, df)

                _format_sheet("Metrics", _round_nearest_100(metrics_df))
                _format_sheet("Forecast", _round_nearest_100(forecast_df) if isinstance(forecast_df, pd.DataFrame) else pd.DataFrame())
                _format_sheet("Income_Annual", income_view)
                _format_sheet("Cashflow_Annual", cash_view)
                _format_sheet("Balance_Annual", balance_view)
                _format_sheet("Plot_Income", income_plot)
                _format_sheet("Plot_Cashflow", cash_plot)
                _format_sheet("Plot_Balance", balance_plot)
                _format_sheet("Plot_Forecast", _round_nearest_100(forecast_df) if isinstance(forecast_df, pd.DataFrame) else pd.DataFrame())

                _add_line_chart("Plot_Income", "Income Statement Trends")
                _add_line_chart("Plot_Cashflow", "Cash Flow Trends")
                _add_line_chart("Plot_Balance", "Balance Sheet Trends")
                _add_line_chart("Plot_Forecast", "Forecast Trends")

                # Ensure all business-plan analysis sections are exported in Excel as dedicated sheets.
                existing_sheet_names = set(writer.sheets.keys())
                for frame_title, frame_df in frames.items():
                    frame_view = _prepare_export_table(frame_df, max_rows=500)
                    if frame_view.empty:
                        continue
                    candidate = "".join(ch for ch in str(frame_title) if ch.isalnum() or ch in (" ", "_", "-"))[:31] or "Analysis"
                    sheet_name = candidate
                    suffix = 1
                    while sheet_name in existing_sheet_names:
                        trimmed = candidate[: max(0, 31 - len(str(suffix)) - 1)]
                        sheet_name = f"{trimmed}_{suffix}"
                        suffix += 1
                    frame_view.to_excel(writer, sheet_name=sheet_name, index=False)
                    existing_sheet_names.add(sheet_name)
                    ws = writer.sheets[sheet_name]
                    ws.set_landscape()
                    ws.fit_to_pages(1, 0)
                    ws.freeze_panes(1, 1)
                    _fit_column_widths_excel(ws, frame_view)
        st.download_button("Download Business Plan Tables (Excel)", data=excel_buf.getvalue(), file_name="Business_Plan_Tables.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        st.download_button("Download Business Plan (PDF)", data=pdf_bytes, file_name="Business_Plan.pdf", mime="application/pdf")
    else:
        st.info("Generate the business plan first to enable downloads.")

    st.markdown("### 7) Additional AI tools to enhance the model")
    st.markdown(
        "- **Anomaly detection** on monthly statements and schedules to flag unusual movements.\n"
        "- **Automatic covenant monitor** (DSCR/LLCR early warning) with threshold alerts.\n"
        "- **Narrative variance analysis** that explains month-on-month and scenario deltas.\n"
        "- **Deal-room pack generator** that bundles assumptions, tables, and charts into investor-ready outputs."
    )

def _render_monte_carlo_page(model: CassavaBioethanolModel, results: Dict[str, object]) -> None:
    st.subheader("Monte Carlo Simulation")

    _ensure_monte_carlo_state()
    st.subheader("Monte Carlo Simulation Configuration")
    current_version = _current_model_version()
    current_scenario = model.scenario
    parameter_source = results.get("input_page_snapshot") if isinstance(results, dict) else None
    if not isinstance(parameter_source, InputLandingPage):
        parameter_source = model.input_page
    parameter_library = _monte_carlo_parameter_library(parameter_source)
    parameter_options = parameter_library["Parameter"].tolist()
    base_value_lookup = (
        parameter_library.set_index("Parameter")["Base Value"].to_dict()
        if not parameter_library.empty
        else {}
    )

    iterations = st.number_input(
        "Iterations",
        min_value=1,
        value=int(st.session_state[MC_ITERATION_STATE_KEY]),
        step=1,
        format="%d",
    )
    st.session_state[MC_ITERATION_STATE_KEY] = int(iterations)

    seed = st.number_input(
        "Random Seed",
        value=int(st.session_state[MC_SEED_STATE_KEY]),
        step=1,
        format="%d",
    )
    st.session_state[MC_SEED_STATE_KEY] = int(seed)

    left_col, right_col = st.columns([1, 2])

    with left_col:
        st.caption(
            "Select which inputs feed the simulation and review their base values from the landing page."
        )

        existing = st.session_state[MC_PARAMETER_STATE_KEY]
        current_parameters = [
            str(value).strip()
            for value in existing.get("Parameter", [])
            if str(value).strip() in parameter_options
        ]
        initial_selection = st.session_state.get(
            MC_SELECTED_PARAMETER_STATE_KEY,
            current_parameters,
        )
        selected_parameters = st.multiselect(
            "Parameters to include",
            options=parameter_options,
            default=initial_selection,
            key="mc_parameter_selector",
        )
        st.session_state[MC_SELECTED_PARAMETER_STATE_KEY] = selected_parameters

        if st.button("Select all variables", use_container_width=True) and parameter_options:
            st.session_state[MC_SELECTED_PARAMETER_STATE_KEY] = parameter_options
            _trigger_rerun()

        if parameter_library.empty:
            st.info("No parameters are available on the landing page to configure the simulation.")
        else:
            st.dataframe(parameter_library, use_container_width=True, hide_index=True)

        custom_parameter = st.text_input(
            "Add custom parameter",
            help="Provide a custom label when you need to sample a variable not listed above.",
        )
        if st.button("Add custom parameter to configuration", use_container_width=True):
            trimmed = custom_parameter.strip()
            if trimmed:
                df = st.session_state[MC_PARAMETER_STATE_KEY]
                if "Parameter" not in df or trimmed not in df["Parameter"].astype(str).tolist():
                    new_row = {
                        column: "" if column in MONTE_CARLO_TEXT_COLUMNS else np.nan
                        for column in MONTE_CARLO_PARAMETER_COLUMNS
                    }
                    new_row["Parameter"] = trimmed
                    new_row["Distribution"] = "Normal"
                    if trimmed in base_value_lookup:
                        new_row["loc"] = base_value_lookup[trimmed]
                    st.session_state[MC_PARAMETER_STATE_KEY] = pd.concat(
                        [df, pd.DataFrame([new_row])], ignore_index=True
                    ).reindex(columns=list(MONTE_CARLO_PARAMETER_COLUMNS))
                selection = st.session_state.get(MC_SELECTED_PARAMETER_STATE_KEY, [])
                if trimmed not in selection:
                    st.session_state[MC_SELECTED_PARAMETER_STATE_KEY] = selection + [trimmed]
                _trigger_rerun()

    with right_col:
        st.caption("Configure distributions and parameters for each selected input.")
        current_params = st.session_state[MC_PARAMETER_STATE_KEY]
        df = current_params.copy()

        if parameter_options:
            retain = set(st.session_state.get(MC_SELECTED_PARAMETER_STATE_KEY, []))
            drop_mask = df["Parameter"].astype(str).isin(set(parameter_options) - retain)
            df = df.loc[~drop_mask].copy()

        selected_parameters = st.session_state.get(MC_SELECTED_PARAMETER_STATE_KEY, [])
        if not selected_parameters:
            st.info("Select one or more parameters to configure the Monte Carlo simulation.")

        updated_rows: List[Dict[str, object]] = []
        distribution_options = available_monte_carlo_distributions()

        for index, parameter in enumerate(selected_parameters):
            existing_rows = df[df["Parameter"].astype(str) == parameter]
            template = {
                column: ("" if column in MONTE_CARLO_TEXT_COLUMNS else np.nan)
                for column in MONTE_CARLO_PARAMETER_COLUMNS
            }
            template["Parameter"] = parameter
            template["Distribution"] = "Normal"

            if not existing_rows.empty:
                row_data = existing_rows.iloc[0].to_dict()
                for column in MONTE_CARLO_PARAMETER_COLUMNS:
                    if column not in row_data:
                        row_data[column] = template[column]
            else:
                row_data = template

            base_value = base_value_lookup.get(parameter)
            slug = re.sub(r"[^0-9A-Za-z]+", "_", parameter).strip("_").lower() or f"param_{index}"

            with st.container():
                st.markdown(f"**{parameter}**")
                if base_value is not None and not (isinstance(base_value, float) and np.isnan(base_value)):
                    st.caption(f"Base value: {base_value}")

                if not distribution_options:
                    st.warning("No probability distributions are available for selection.")
                    row_data["Distribution"] = ""
                    spec = None
                else:
                    current_distribution = str(row_data.get("Distribution", "") or "Normal")
                    if current_distribution not in distribution_options:
                        current_distribution = distribution_options[0]

                    distribution = st.selectbox(
                        "Distribution",
                        options=distribution_options,
                        index=distribution_options.index(current_distribution),
                        key=f"mc_dist_{slug}",
                        help="Probability distribution used for sampling during the Monte Carlo run.",
                    )
                    row_data["Distribution"] = distribution
                    spec = MONTE_CARLO_DISTRIBUTIONS.get(distribution)
                relevant_fields: List[str] = []
                if spec:
                    relevant_fields.extend(spec.shape_params)
                    relevant_fields.extend(spec.keyword_params)

                if not relevant_fields:
                    st.caption("No additional parameters required for this distribution.")
                else:
                    num_columns = min(3, max(1, len(relevant_fields)))
                    field_columns = st.columns(num_columns)
                    for field_index, field_name in enumerate(relevant_fields):
                        field_column = field_columns[field_index % num_columns]
                        raw_value = row_data.get(field_name)
                        placeholder = ""
                        if field_name == "loc" and base_value is not None and not (
                            isinstance(base_value, float) and np.isnan(base_value)
                        ):
                            placeholder = str(base_value)

                        key_suffix = f"{slug}_{field_name}"
                        if field_name == "pvals":
                            default_value = "" if pd.isna(raw_value) else str(raw_value)
                            value = field_column.text_input(
                                field_name,
                                value=default_value,
                                key=f"mc_param_{key_suffix}",
                                help="Comma-separated probabilities (e.g. 0.2,0.3,0.5) for the Multinomial distribution.",
                            )
                            row_data[field_name] = value
                        else:
                            default_value = "" if pd.isna(raw_value) else str(raw_value)
                            value = field_column.text_input(
                                field_name,
                                value=default_value,
                                key=f"mc_param_{key_suffix}",
                                placeholder=placeholder,
                            )
                            row_data[field_name] = value

                for column in MONTE_CARLO_PARAMETER_COLUMNS:
                    if column in {"Parameter", "Distribution"}:
                        continue
                    if spec and column in relevant_fields:
                        continue
                    row_data[column] = "" if column in MONTE_CARLO_TEXT_COLUMNS else np.nan

                updated_rows.append(row_data)

                st.divider()

        edited_params = pd.DataFrame(updated_rows, columns=list(MONTE_CARLO_PARAMETER_COLUMNS))
        edited_params = edited_params.reindex(columns=list(MONTE_CARLO_PARAMETER_COLUMNS))

        numeric_columns = [
            column for column in MONTE_CARLO_PARAMETER_COLUMNS if column not in MONTE_CARLO_TEXT_COLUMNS
        ]
        for column in numeric_columns:
            if column in edited_params:
                edited_params[column] = pd.to_numeric(edited_params[column], errors="coerce")
        for column in MONTE_CARLO_TEXT_COLUMNS:
            if column in edited_params:
                edited_params[column] = edited_params[column].astype("string").fillna("").astype(object)
        st.session_state[MC_PARAMETER_STATE_KEY] = edited_params

        with st.expander("Distribution reference", expanded=False):
            st.dataframe(
                _monte_carlo_distribution_table(),
                use_container_width=True,
                hide_index=True,
            )

    config_signature = _monte_carlo_signature(edited_params, int(iterations), int(seed))

    cache: Dict[str, Dict[str, object]] = st.session_state.setdefault(MC_CACHE_KEY, {})
    cached_entry = cache.get(current_scenario)
    if cached_entry and (
        cached_entry.get("version") != current_version
        or cached_entry.get("signature") != config_signature
    ):
        cache.pop(current_scenario, None)
        st.session_state[MC_CACHE_KEY] = cache
        cached_entry = None

    run_requested = st.button(
        "Run Monte Carlo Simulation",
        key=f"mc_run_{current_scenario.lower()}",
    )

    if run_requested:
        base_page = copy.deepcopy(results.get("input_page_snapshot", model.input_page))
        with st.spinner("Running Monte Carlo simulation..."):
            mc_model = CassavaBioethanolModel(copy.deepcopy(base_page))
            mc_model.scenario = current_scenario
            mc_results = monte_carlo_simulation(
                mc_model,
                parameter_configs=edited_params,
                iterations=int(iterations),
                random_seed=int(seed),
            )
        cache[current_scenario] = {
            "version": current_version,
            "signature": config_signature,
            "results": mc_results,
        }
        st.session_state[MC_CACHE_KEY] = cache
        cached_entry = cache[current_scenario]
    elif cached_entry is None:
        st.info("Monte Carlo results are not available. Click 'Run Monte Carlo Simulation' to generate them.")
        return

    mc_results = cached_entry["results"]
    if mc_results.empty:
        st.info("Monte Carlo results are not available for the current configuration.")
        return

    st.subheader("Summary Statistics")
    summary = mc_results.describe(percentiles=[0.1, 0.25, 0.5, 0.75, 0.9]).T.reset_index()
    summary = summary.rename(columns={"index": "Metric"})
    stat_columns = [
        column
        for column in summary.columns
        if column != "Metric" and summary[column].notna().any()
    ]
    if stat_columns:
        summary_melt = summary.melt(
            id_vars="Metric",
            value_vars=stat_columns,
            var_name="Statistic",
            value_name="Value",
        )
        summary_chart = px.bar(
            summary_melt,
            x="Statistic",
            y="Value",
            color="Statistic",
            facet_col="Metric",
            facet_col_wrap=2,
            title="Monte Carlo summary statistics by metric",
        )
        summary_chart.update_layout(showlegend=False)
        st.plotly_chart(summary_chart, use_container_width=True)
    else:
        st.info("Summary statistics are not available for the current simulation results.")

    numeric_columns = mc_results.select_dtypes(include=[np.number]).columns.tolist()
    if numeric_columns:
        st.subheader("Distribution Visualisations")
        for column in numeric_columns:
            distribution_chart = px.histogram(
                mc_results,
                x=column,
                nbins=min(50, max(10, int(np.sqrt(len(mc_results))))),
                marginal="box",
                title=f"Distribution for {column}",
            )
            st.plotly_chart(distribution_chart, use_container_width=True)
    else:
        st.info("Monte Carlo results do not contain numeric metrics to plot.")

def main() -> None:
    st.title("Cassava_Bioethanol Financial Model")
    st.caption("Adjust the assumptions, run the project finance model, and inspect the outputs across dedicated dashboards.")

    input_page = _load_session_inputs()
    _sync_projection_from_session(input_page)
    scenario_options = list(CassavaBioethanolModel.SCENARIOS)
    if "selected_scenario" not in st.session_state:
        st.session_state.selected_scenario = scenario_options[0]

    action_cols = st.columns([1, 1, 1])
    with action_cols[0]:
        recalc = st.button("Recalculate model", type="primary")
    with action_cols[1]:
        scenario_index = scenario_options.index(st.session_state.selected_scenario)
        selected_choice = st.selectbox(
            "Scenario",
            scenario_options,
            index=scenario_index,
            key="scenario_select",
        )
    download_container = action_cols[2].container()

    if selected_choice != st.session_state.selected_scenario:
        st.session_state.selected_scenario = selected_choice
        st.session_state.scenario_payloads = {}
        st.session_state.excel_bytes_map = {}
        st.session_state[MC_CACHE_KEY] = {}
        st.session_state.pop("mc_cache", None)
        st.session_state.pop("mc_cache_version", None)
        st.session_state.pop("mc_cache_scenario", None)
        st.session_state[SENSITIVITY_CACHE_KEY] = {}
        st.session_state[SCENARIO_CACHE_KEY] = {}

    selected_scenario = st.session_state.selected_scenario

    st.caption("Use the navigation tabs to move between the input landing page and the analytical dashboards.")

    tabs = st.tabs(
        [
            "Input Landing Page",
            "Key Metrics Dashboard",
            "Financial Performance",
            "Financial Position",
            "Cash Flow Statement",
            "Sensitivity Analyses",
            "Scenario / IFs Analysis",
            "Monte Carlo Simulation",
            "RAG Assistant",
        ]
    )

    with tabs[0]:
        st.subheader("Input Landing Page")
        st.info("Edit the assumptions and press 'Recalculate model' to refresh the financial outputs.")
        _update_projection(input_page)
        _key_assumptions_controls(input_page.global_inputs)
        _modify_default_inputs(input_page)
        _render_production_panel(input_page)
        _sync_table_editors(input_page)
        _apply_dependent_updates(input_page)
        _editable_tables(input_page)

    snapshot = st.session_state.get("input_snapshot")
    snapshot_projection = None
    if snapshot is not None:
        snapshot_projection = (
            int(snapshot.projection.start_year),
            int(snapshot.projection.end_year),
            str(snapshot.projection.planning_start),
        )
    current_projection = (
        int(input_page.projection.start_year),
        int(input_page.projection.end_year),
        str(input_page.projection.planning_start),
    )
    projection_changed = snapshot_projection is not None and snapshot_projection != current_projection
    inputs_dirty = st.session_state.get("inputs_dirty", False)

    if (
        recalc
        or projection_changed
        or inputs_dirty
        or "scenario_payloads" not in st.session_state
    ):
        st.session_state.scenario_payloads = {}
        st.session_state.excel_bytes_map = {}
        st.session_state.input_snapshot = copy.deepcopy(input_page)
        _bump_model_version()
        st.session_state[MC_CACHE_KEY] = {}
        st.session_state.pop("mc_cache", None)
        st.session_state.pop("mc_cache_version", None)
        st.session_state.pop("mc_cache_scenario", None)
        st.session_state[SENSITIVITY_CACHE_KEY] = {}
        st.session_state[SCENARIO_CACHE_KEY] = {}
        st.session_state.inputs_dirty = False

    snapshot = st.session_state.get("input_snapshot")
    if snapshot is None:
        snapshot = copy.deepcopy(input_page)
        st.session_state.input_snapshot = snapshot

    model, results = _ensure_scenario_payload(selected_scenario, snapshot)
    st.session_state.model_results = (model, results)

    excel_map: Dict[str, bytes] = st.session_state.setdefault("excel_bytes_map", {})
    excel_bytes = excel_map.get(selected_scenario)

    model.scenario = selected_scenario

    with download_container:
        if not excel_bytes:
            if st.button("Prepare Excel Model", key=f"prepare_excel_{selected_scenario.lower()}"):
                with st.spinner("Preparing Excel workbook..."):
                    excel_bytes = _generate_excel_bytes(model, results, selected_scenario)
                excel_map[selected_scenario] = excel_bytes
                st.session_state.excel_bytes_map = excel_map
        if excel_bytes:
            st.download_button(
                "Download Excel Model",
                data=excel_bytes,
                file_name="Cassava_Bioethanol_Financial_Model.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            if st.button(
                "Clear Prepared Excel",
                key=f"clear_excel_{selected_scenario.lower()}",
            ):
                excel_map.pop(selected_scenario, None)
                st.session_state.excel_bytes_map = excel_map
                excel_bytes = None
        if not excel_bytes:
            st.info("Click 'Prepare Excel Model' to generate the workbook for download.")

    with tabs[1]:
        _render_key_metrics(model, results)

    with tabs[2]:
        _render_financial_performance(results)

    with tabs[3]:
        _render_financial_position(results)

    with tabs[4]:
        _render_cash_flow_page(results)

    with tabs[5]:
        _render_sensitivity_page(model, results)

    with tabs[6]:
        _render_scenario_page(model, results)

    with tabs[7]:
        _render_monte_carlo_page(model, results)

    with tabs[8]:
        _render_rag_assistant_page(model, results)


if __name__ == "__main__":
    main()
